from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"D:\长剧情")
DELIVER = ROOT / "软著材料" / "交付" / "2026-08-08_织幕V1.0"
QA = ROOT / "软著材料" / "渲染检查" / "2026-08-08_织幕V1.0"
SOURCE_TEXT = ROOT / "软著材料" / "输出" / "源代码交存稿_织幕V1.0.txt"


def document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main() -> None:
    source_text = SOURCE_TEXT.read_text(encoding="utf-8")
    source_file_match = re.search(r"^纳入文件数：(\d+)$", source_text, flags=re.MULTILINE)
    source_line_match = re.search(r"^源代码总行数：(\d+)$", source_text, flags=re.MULTILINE)
    assert source_file_match and source_line_match, "source statistics are missing"
    source_file_count = int(source_file_match.group(1))
    source_line_count = int(source_line_match.group(1))

    docx_files = sorted(DELIVER.glob("*.docx"))
    assert len(docx_files) == 4, f"expected 4 docx files, got {len(docx_files)}"
    docs = {path.name: Document(path) for path in docx_files}
    texts = {name: document_text(doc) for name, doc in docs.items()}

    all_text = "\n".join(texts.values())
    required = [
        "织幕长线剧本杀自动化叙事与运营平台软件",
        "V1.0",
        "用户自备 API Key",
        "平台自有 API",
        "DeepSeek",
        "OpenRouter",
        "SiliconFlow",
        "账号分级",
        "活态故事总览",
        "语义大纲",
        "六位验证码",
        "三端用户资料",
        "资产配额",
        "机制设计工作台",
        "玩家决策提交",
        "权威结算",
        "秘密承诺",
        "服务器轮次时钟",
        "期望修订号",
    ]
    for phrase in required:
        assert phrase in all_text, f"missing required phrase: {phrase}"
    assert "productionTrust 7/7" not in all_text
    assert "平台积分、付费套餐和平台自有 API 额度" in all_text
    assert f"{source_file_count} 个自主项目源文件" in all_text
    assert f"{source_line_count:,} 行" in all_text
    assert "建议：已发表" in all_text
    assert "建议：2026 年 8 月 7 日" in all_text

    manual_name = next(name for name in texts if name.startswith("02_"))
    assert len(docs[manual_name].inline_shapes) == 4

    source_name = next(name for name in texts if name.startswith("03_"))
    source_doc = docs[source_name]
    code_paragraphs = [
        paragraph.text.splitlines()
        for paragraph in source_doc.paragraphs
        if paragraph.text.startswith("01:")
    ]
    assert len(code_paragraphs) == 60
    assert all(len(lines) == 50 for lines in code_paragraphs)
    assert all(lines[-1].startswith("50:") for lines in code_paragraphs)

    assert f"纳入文件数：{source_file_count}" in source_text
    assert f"源代码总行数：{source_line_count}" in source_text
    numbered_lines = [line for line in source_text.splitlines() if re.match(r"^\d{2}:", line)]
    assert len(numbered_lines) == 3000
    assert [line for page in code_paragraphs for line in page] == numbered_lines, "DOCX source lines are not continuous with generated source text"
    assert "连续性说明：前 30 页与后 30 页分别从确定性文件序列连续截取" in source_text
    for required_source in [
        "shared/mechanism-design.js",
        "src/views/creator-mechanism-workbench.js",
        "backend/src/room-mechanism-submission-service.js",
        "play/src/runtime/game-action-controller.js",
        "host/src/views/host-mechanism-workspace.js",
        "backend/src/routes/host-mechanism-runtime-routes.js",
    ]:
        assert required_source in source_text, f"missing deposited feature evidence: {required_source}"
    secret_patterns = {
        "private_key": r"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----",
        "github_token": r"\bgh[pousr]_[A-Za-z0-9]{30,}\b",
        "openai_style_key": r"\bsk-[A-Za-z0-9_-]{20,}\b",
        "database_url": r"\bpostgres(?:ql)?://[^\s\"']+",
        "jwt": r"\beyJ[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\b",
    }
    secret_hits = {
        name: len(re.findall(pattern, source_text, flags=re.IGNORECASE))
        for name, pattern in secret_patterns.items()
    }
    assert not any(secret_hits.values()), secret_hits

    pdf_pages = {}
    for directory in sorted(path for path in QA.iterdir() if path.is_dir()):
        pdfs = list(directory.glob("*.pdf"))
        assert len(pdfs) == 1, f"{directory.name}: expected one PDF"
        pdf_pages[directory.name] = len(PdfReader(pdfs[0]).pages)
    source_pdf_name = next(name for name in pdf_pages if name.startswith("03_"))
    assert pdf_pages[source_pdf_name] == 60
    assert pdf_pages[next(name for name in pdf_pages if name.startswith("02_"))] >= 10

    delivery_pdfs = sorted(DELIVER.glob("*.pdf"))
    assert len(delivery_pdfs) == 4, f"expected 4 delivery PDFs, got {len(delivery_pdfs)}"
    delivery_pdf_pages = {path.stem: len(PdfReader(path).pages) for path in delivery_pdfs}
    assert delivery_pdf_pages == pdf_pages, "delivery PDF page counts differ from QA PDFs"

    evidence_path = DELIVER / "05_织幕V1.0_相对上一版新增功能证据.txt"
    evidence_text = evidence_path.read_text(encoding="utf-8")
    for phrase in ["2026-08-06_织幕V1.0", "机制设计工作台", "玩家决策提交", "主持端权威结算", "跨端同步与隐私", "源程序交存稿证据页"]:
        assert phrase in evidence_text, f"evidence report missing: {phrase}"

    result = {
        "docx_files": [path.name for path in docx_files],
        "docx_opened": len(docs),
        "manual_images": len(docs[manual_name].inline_shapes),
        "source_pages": len(code_paragraphs),
        "source_lines_per_page": sorted({len(lines) for lines in code_paragraphs}),
        "source_security_hits": secret_hits,
        "pdf_pages": pdf_pages,
        "delivery_pdf_pages": delivery_pdf_pages,
        "evidence_report": evidence_path.name,
        "placeholder_count": all_text.count("待填写") + all_text.count("待确认"),
        "status": "PASS",
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
