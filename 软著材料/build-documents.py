from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\长剧情")
MATERIALS = ROOT / "软著材料"
PACKAGE_DATE = "2026-08-08"
PACKAGE_DATE_CN = "2026 年 8 月 8 日"
OUTPUT = MATERIALS / "交付" / f"{PACKAGE_DATE}_织幕V1.0"
SOURCE_TEXT = MATERIALS / "输出" / "源代码交存稿_织幕V1.0.txt"
PREVIOUS_PACKAGE = "2026-08-06_织幕V1.0"
FEATURE_IMPLEMENTED_DATE_CN = "2026 年 8 月 7 日"
FEATURE_EVIDENCE = {
    "机制设计工作台": [
        "shared/mechanism-design.js",
        "shared/mechanism-interactions.js",
        "src/views/creator-mechanism-workbench.js",
        "src/runtime/actions-creator-cockpit.js",
    ],
    "玩家决策提交": [
        "backend/migrations/107_room_mechanism_decision_submissions.sql",
        "backend/migrations/108_room_mechanism_round_clock.sql",
        "backend/src/room-mechanism-submission-service.js",
        "backend/src/routes/player-progress-routes.js",
        "play/src/runtime/game-action-controller.js",
        "play/src/views/game-home-views.js",
    ],
    "主持端权威结算": [
        "backend/src/room-mechanism-runtime-service.js",
        "backend/src/routes/host-mechanism-runtime-routes.js",
        "host/src/runtime/host-mechanism-controller.js",
        "host/src/views/host-mechanism-workspace.js",
    ],
}
FEATURE_TEST_EVIDENCE = [
    "scripts/mechanism-design.test.mjs",
    "scripts/creator-mechanism-workbench.test.mjs",
    "play/test/mechanism-progress.test.mjs",
    "host/test/mechanism-workspace.test.mjs",
    "backend/test/mechanism-submission-service.test.js",
    "backend/test/mechanism-runtime.test.js",
    "backend/test/room-event-audience.test.js",
]
SCREENSHOTS = {
    "creator": ROOT / "site" / "public" / "assets" / "zhimu-screenshot-creator.png",
    "host": ROOT / "site" / "public" / "assets" / "zhimu-screenshot-host.png",
    "play": ROOT / "site" / "public" / "assets" / "zhimu-screenshot-play.png",
    "archive": ROOT / "site" / "public" / "assets" / "zhimu-screenshot-archive.png",
}

SOFTWARE_NAME = "织幕长线剧本杀自动化叙事与运营平台软件"
SHORT_NAME = "织幕"
VERSION = "V1.0"

# compact_reference_guide preset with explicit A4/CJK overrides required by
# software-copyright submission conventions.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "243443"
MUTED = "687887"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FB"
LIGHT_GRAY = "F1F3F5"
MID_GRAY = "D6DEE5"
YELLOW = "FFF4CE"
YELLOW_BORDER = "E6B800"
WHITE = "FFFFFF"
PAGE_WIDTH_TWIPS = 11906  # A4
BODY_WIDTH_TWIPS = 9660


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: Sequence[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_repeatable_metadata(doc: Document, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "织幕软件著作权材料"
    props.keywords = "织幕, 软件著作权, V1.0"
    props.comments = "根据当前项目实现整理；申请人身份与权属信息需提交前确认。"
    props.last_modified_by = "织幕软件著作权材料"


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None,
                 bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 26, "203748", 0, 12),
        ("Subtitle", 11, MUTED, 0, 10),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if "Heading" in style_name or style_name == "Title":
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Microsoft YaHei"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def set_page_layout(doc: Document, top=0.70, bottom=0.70, left=0.78, right=0.78) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)


def add_field(paragraph, instruction: str, result: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def add_running_header_footer(doc: Document, short_title: str) -> None:
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(short_title)
        set_run_font(hr, size=8, color=MUTED)
        hp.paragraph_format.space_after = Pt(0)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = fp.add_run("— ")
        set_run_font(prefix, size=8, color=MUTED)
        page_run = add_field(fp, "PAGE", "1")
        set_run_font(page_run, size=8, color=MUTED)
        suffix = fp.add_run(" —")
        set_run_font(suffix, size=8, color=MUTED)
        fp.paragraph_format.space_before = Pt(0)


def new_document(title: str, subject: str, short_header: str) -> Document:
    doc = Document()
    configure_styles(doc)
    set_page_layout(doc)
    set_repeatable_metadata(doc, title, subject)
    add_running_header_footer(doc, short_header)
    return doc


def add_title_page(doc: Document, title: str, subtitle: str,
                   kicker: str = "SOFTWARE COPYRIGHT MATERIALS") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(kicker)
    set_run_font(r, size=9, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(title)
    set_run_font(r, size=26, bold=True, color="203748")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=11, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{SOFTWARE_NAME}\n{VERSION}")
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r = p.add_run(f"材料编制日期：{PACKAGE_DATE_CN}")
    set_run_font(r, size=9.5, color=MUTED)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=TEXT)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, color=TEXT)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.widow_control = True
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def add_steps(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.widow_control = True
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def add_notice(doc: Document, title: str, body: str, *, warning: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [BODY_WIDTH_TWIPS])
    cell = table.cell(0, 0)
    set_cell_shading(cell, YELLOW if warning else PALE_BLUE)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), YELLOW_BORDER if warning else BLUE)
    borders.append(left)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=9.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], header_fill: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=9.5, bold=True, color=WHITE if header_fill == BLUE else TEXT)

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=9.3, color=TEXT)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, image_path: Path, caption: str, figure_no: int) -> None:
    if not image_path.exists():
        add_notice(doc, "截图缺失", f"未找到 {image_path.name}，提交前请补充。", warning=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.22))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    r = cap.add_run(f"图 {figure_no}  {caption}")
    set_run_font(r, size=9, color=MUTED)


def add_manual_index(doc: Document, entries: Sequence[tuple[str, str]]) -> None:
    doc.add_heading("文档导读", level=1)
    add_body(doc, "本说明书按“软件概述—运行环境—功能模块—典型流程—安全与维护”的顺序说明织幕 V1.0 的实际功能。")
    add_table(doc, ["章节", "内容"], entries, [1850, 7810], header_fill=BLUE)


def source_stats() -> tuple[int, int]:
    text = SOURCE_TEXT.read_text(encoding="utf-8")
    file_match = re.search(r"^纳入文件数：(\d+)$", text, flags=re.MULTILINE)
    line_match = re.search(r"^源代码总行数：(\d+)$", text, flags=re.MULTILINE)
    if not file_match or not line_match:
        raise RuntimeError("源程序交存稿缺少文件数或总行数统计，请先重新运行 prepare-source.mjs")
    return int(file_match.group(1)), int(line_match.group(1))


def source_evidence_pages() -> dict[str, list[int]]:
    """Map deposited evidence file markers to their rendered source pages."""
    numbered = [
        line[4:] if len(line) >= 4 else ""
        for line in SOURCE_TEXT.read_text(encoding="utf-8").splitlines()
        if re.match(r"^\d{2}:", line)
    ]
    pages: dict[str, list[int]] = {}
    for index, line in enumerate(numbered):
        match = re.match(r"^// ===== 文件：(.+) =====$", line)
        if not match:
            continue
        rel = match.group(1)
        pages.setdefault(rel, []).append(index // 50 + 1)
    return pages


def evidence_page_label(paths: Sequence[str]) -> str:
    page_map = source_evidence_pages()
    pages = sorted({page for path in paths for page in page_map.get(path, [])})
    return "、".join(f"第 {page} 页" for page in pages) if pages else "代码仓库与测试证据"


def build_application_info() -> Path:
    path = OUTPUT / "01_织幕软件著作权登记信息表_V1.0.docx"
    source_file_count, source_line_count = source_stats()
    doc = new_document(
        "织幕软件著作权登记信息表 V1.0",
        "用于填写中国版权保护中心软件著作权登记申请表的资料底稿",
        "织幕 V1.0 · 登记信息表",
    )
    add_title_page(
        doc,
        "软件著作权登记信息表",
        "申请资料底稿 · 可直接据此填写官方在线申请表",
    )
    add_notice(
        doc,
        "提交前必须确认",
        "带“【待填写】”或“【待确认】”的内容涉及申请主体、完成日期、发表状态和权属，不能由项目代码推断。正式申请表、说明书和源程序中的软件全称、简称、版本号必须完全一致。",
        warning=True,
    )

    add_heading(doc, "一、软件基本信息", 1)
    rows = [
        ("软件全称", SOFTWARE_NAME, "建议保持不变"),
        ("软件简称", SHORT_NAME, "建议保持不变"),
        ("版本号", VERSION, "首次登记版本"),
        ("软件类别", "应用软件；网络平台软件", "按申请系统可选项填写"),
        ("软件作品说明", "原创软件", "【待确认】如含合作、委托或受让请改填"),
        ("开发方式", "独立开发", "【待确认】"),
        ("权利取得方式", "原始取得", "【待确认】"),
        ("权利范围", "全部权利", "【待确认】"),
        ("开发完成日期", "建议：2026 年 8 月 7 日", "按本轮机制功能完成提交与当前 V1.0 冻结点整理；申请人确认"),
        ("发表状态", "建议：已发表", "申请人已说明网站可由公众访问；若当时仅开放宣传页则改按实际情况填写"),
        ("首次发表日期", "【待填写：V1.0 首次向公众提供软件功能的日期】", "不得凭提交日期或代码提交日期猜填"),
        ("首次发表地点", "中国 / 互联网", "申请人确认"),
    ]
    add_table(doc, ["项目", "建议填写", "说明"], rows, [2100, 4300, 3260])

    add_heading(doc, "二、申请人及联系人信息", 1)
    rows = [
        ("申请主体类型", "【个人 / 企业 / 其他组织】"),
        ("著作权人名称", "【待填写，与证件名称完全一致】"),
        ("证件类型", "【居民身份证 / 营业执照等】"),
        ("证件号码", "【待填写】"),
        ("国籍或注册地", "【待填写】"),
        ("联系地址", "【待填写】"),
        ("邮政编码", "【待填写】"),
        ("联系人", "【待填写】"),
        ("手机号码", "【待填写】"),
        ("电子邮箱", "【待填写】"),
    ]
    add_table(doc, ["项目", "填写内容"], rows, [2600, 7060])

    add_heading(doc, "三、开发与运行环境", 1)
    rows = [
        ("开发硬件环境", "通用 x86-64 个人计算机；建议 8 GB 以上内存、20 GB 以上可用存储空间"),
        ("运行硬件环境", "云服务器或容器运行环境；客户端为可访问互联网的桌面或移动终端"),
        ("开发操作系统", "Windows 11；Node.js 24.13.x 开发运行环境"),
        ("软件运行平台", "服务器端 Node.js 24.13.x、PostgreSQL 17；客户端为现代浏览器"),
        ("开发工具", "Visual Studio Code、Git、npm、Vite、数据库迁移工具"),
        ("开发语言", "JavaScript、SQL、HTML、CSS"),
        ("后端技术", "Fastify 5、PostgreSQL 17、REST API、SSE"),
        ("前端技术", "Vite 8、原生 ES Modules、HTML5、CSS3"),
        ("对象存储", "兼容 S3 API 的对象存储（生产环境采用 Cloudflare R2）"),
        ("支撑软件", "现代 Chromium、Firefox 或 Safari 浏览器；可选监控与恶意文件扫描服务"),
    ]
    add_table(doc, ["项目", "建议填写"], rows, [2600, 7060])

    add_heading(doc, "四、软件功能说明（申请表摘要）", 1)
    add_body(
        doc,
        "织幕是一套面向线上长线剧本杀、跑团和互动叙事的创作、主持与玩家协作平台。软件提供创作者端、主持端、玩家端和官网入口，覆盖世界、角色与章节创建，创作宪法、真相链和故事图谱整理，活态故事总览与作者确认，机制设计工作台，语义大纲生成及质量校验，场景、线索、物品与结构化规则编排，作品诊断与 AI 玩家测试，运行房创建、角色分配、玩家机制决策提交、主持端权威结算与超时默认方案、实时事件同步、检查点恢复、复盘归档、资产配额与安全上传、三端独立用户资料、内测账号分级及审计等功能。用户侧 AI 功能采用用户自备 API Key（BYOK）模式，可配置多家模型服务；平台积分、付费套餐和平台自有 API 额度暂不作为 V1.0 用户功能开放。",
    )

    add_heading(doc, "五、技术特点说明（申请表摘要）", 1)
    add_bullets(doc, [
        "采用 Creator、Host、Player 和官网分端架构，各端按创作者、主持人、玩家和公开访客的角色职责独立部署。",
        "后端基于 Node.js 24.13.x 与 Fastify 5，使用 PostgreSQL 17 作为业务数据真相源，以结构化迁移管理数据模型。",
        "通过 REST API 与服务器发送事件（SSE）提供统一业务访问和实时同步，并结合持久事件日志、outbox/journal、PostgreSQL LISTEN/NOTIFY 与轮询补偿实现断线恢复。",
        "创作数据按世界、角色、章节、场景、线索、物品、调查点、规则和运行房等领域对象组织，支持从创作模板到房间实例的版本边界。",
        "活态故事总览可综合世界简介、创作宪法、角色、关系、核心事实与章节生成故事脊柱候选，区分作者已确认、AI 暂拟和待决定内容，并保留来源与修订信息。",
        "语义大纲流水线对玩家身份、章节因果、题材资源、证据、结局与角色责任进行结构化校验，减少跨模块内容漂移。",
        "结构化规则引擎仅接受受控条件和动作，不允许用户提交任意 JavaScript 执行。",
        "机制设计工作台把作者机制七问、线上表现、玩家提示、主持提示、提交模式、截止时间和默认方案保存为结构化合同；草稿与作者已确认状态明确分离。",
        "玩家提交只表达公开倾向或秘密承诺，不直接修改权威运行态；主持端查看提交汇总后作最终确认，服务端校验主持权限、作品绑定、期望修订号和服务器轮次时钟。",
        "限时机制到期前禁止提前执行默认后果，到期后阻止普通结算并要求使用作者预设默认方案；玩家投影隐藏作者内部机制键，秘密承诺详情仅主持端可见。",
        "用户侧模型连接支持 DeepSeek、OpenAI、OpenRouter、通义千问、智谱 GLM、SiliconFlow 及自定义 OpenAI 兼容接口；API Key 使用 AES-256-GCM 加密保存，运行时仅按用户授权调用。",
        "安全能力包含后端权限校验、HttpOnly 会话、内容安全策略、限流、SSRF 防护、上传扫描、结构化审计和账号数据导出/删除流程。",
        "账号注册支持六位邮箱验证码及邮件一键验证；Creator、Host、Player 可分别维护显示名称与头像，避免不同使用身份互相覆盖。",
        "对象存储、健康检查、指标、链路追踪、告警、备份与回滚流程为生产部署和维护提供支撑。",
    ])

    add_heading(doc, "六、源程序与文档规模", 1)
    add_table(doc, ["项目", "当前材料"], [
        ("源程序鉴别材料", "普通交存：前 30 页 + 后 30 页，共 60 页，每页 50 行"),
        ("源程序纳入范围", "Creator、Host、Player、共享模块、后端业务代码及数据库迁移"),
        ("本次源代码统计", f"{source_file_count} 个自主项目源文件，约 {source_line_count:,} 行；不包含依赖、构建产物、测试、环境文件和密钥"),
        ("文档鉴别材料", "《织幕长线剧本杀自动化叙事与运营平台软件 V1.0 操作说明书》"),
        ("文字与截图", "以当前 V1.0 功能和实际产品截图编制"),
    ], [2600, 7060])

    add_heading(doc, "七、权属情况核对", 1)
    add_notice(
        doc,
        "权属风险提示",
        "如果存在共同开发、委托开发、职务开发、代码受让、外包交付或开源代码深度改造，应按真实情况准备合同、任务书、劳动关系、转让或许可证明。不要仅因代码仓库由你控制就直接勾选“独立开发、全部权利”。",
        warning=True,
    )
    add_table(doc, ["核对项", "当前建议", "提交前确认"], [
        ("是否合作开发", "否", "□ 确认  □ 修改"),
        ("是否委托开发", "否", "□ 确认  □ 修改"),
        ("是否属于职务开发", "【待确认】", "□ 个人独立  □ 公司职务"),
        ("是否存在权利转让", "否", "□ 确认  □ 修改"),
        ("是否包含第三方商业源码", "【待确认】", "□ 无  □ 有并已获授权"),
        ("开源依赖使用", "有，作为运行依赖，不作为本次自主源程序交存主体", "□ 已核对许可证"),
    ], [3000, 3400, 3260])

    add_heading(doc, "八、提交前签字确认", 1)
    add_body(doc, "本人/本单位确认：以上软件名称、版本、开发完成日期、发表状态、申请主体和权属情况真实、准确，并与正式申请表、源程序鉴别材料及文档鉴别材料保持一致。")
    add_table(doc, ["项目", "签署"], [
        ("申请人或经办人", "____________________________"),
        ("签字/盖章", "____________________________"),
        ("确认日期", "________ 年 ____ 月 ____ 日"),
    ], [3200, 6460])

    doc.save(path)
    return path


def build_manual() -> Path:
    path = OUTPUT / "02_织幕软件操作说明书_V1.0.docx"
    doc = new_document(
        f"{SOFTWARE_NAME} {VERSION} 操作说明书",
        "软件著作权登记文档鉴别材料",
        "织幕 V1.0 · 操作说明书",
    )
    add_title_page(
        doc,
        "软件操作说明书",
        "软件著作权登记文档鉴别材料",
    )
    add_notice(
        doc,
        "版本边界",
        "本说明书描述织幕 V1.0 当前已实现的创作、活态故事总览、机制设计工作台、语义大纲、玩家决策提交、主持端权威结算、三端用户资料及用户自备模型连接功能。平台积分、付费套餐和平台自有 API 额度已作为后续能力预留，但在本版本内不对普通用户开放。",
    )
    add_manual_index(doc, [
        ("第 1—3 章", "软件概述、适用范围、运行与部署环境"),
        ("第 4 章", "系统组成、数据对象与三端协作关系"),
        ("第 5—8 章", "账号与模型连接、创作者端、主持端、玩家端操作"),
        ("第 9—11 章", "实时同步、存档复盘、资产与运营管理"),
        ("第 12—15 章", "典型流程、异常处理、安全、维护与版本说明"),
    ])

    add_heading(doc, "1. 软件概述", 1)
    add_body(doc, f"{SOFTWARE_NAME}（简称“{SHORT_NAME}”）是一套面向线上长线剧本杀、跑团和互动叙事场景的软件平台。系统围绕内容创作、结构编排、测试诊断、运行主持、玩家互动与复盘归档建立一套连续工作流。")
    add_body(doc, "软件将创作阶段的世界模板与实际开局后的房间实例分离：创作者维护可复用的角色、章节、线索和规则；开局时系统从模板创建独立房间，由主持人推进，并仅向每位玩家展示其被授权看到的内容。")
    add_bullets(doc, [
        "创作者：建立世界、角色和章节，整理真相链与伏笔，使用机制设计工作台定义玩家反复行动和跨端结算合同，进行作品诊断与 AI 玩家测试。",
        "主持人：选择房间，检查玩家状态，分配或调整角色，推进阶段，查看玩家机制提交汇总，并对唯一权威运行态执行最终结算。",
        "玩家：通过邀请码进入房间，查看个人角色与章节内容，调查线索，提交机制倾向或秘密承诺，参与投票与私密行动，并在结束后查看复盘。",
        "运营与内测管理：处理内测申请、账号分级、内容审核、反馈、审计和系统健康状态。",
    ])

    add_heading(doc, "2. 适用范围与术语", 1)
    add_table(doc, ["术语", "说明"], [
        ("世界", "一部可复用的互动叙事作品模板，包含角色、章节、线索、规则等内容。"),
        ("运行房", "由世界模板创建的一次独立开局实例，保存成员、进度、事件和复盘数据。"),
        ("Creator", "创作者端，用于内容生产、测试、发布准备和账号管理。"),
        ("Host", "主持端，用于现场推进和运行控制。"),
        ("Player", "玩家端，用于加入房间、阅读、调查和互动。"),
        ("BYOK", "Bring Your Own Key，用户自行提供模型服务商 API Key 并消耗自己的额度。"),
        ("检查点", "运行房在特定时刻保存的可恢复状态。"),
        ("SSE", "服务器发送事件，用于将运行状态和房间事件实时推送到浏览器。"),
        ("机制包", "由作者设计并随作品版本发布的结构化轮次、决策、资源、证据、结局与交互合同。"),
        ("玩家提交", "玩家对当前机制选项表达的公开倾向或秘密承诺，可在截止前更新，不等同于最终结算。"),
        ("权威结算", "主持端发起、服务端校验权限与修订后写入房间唯一运行态的最终机制操作。"),
    ], [2400, 7260])

    add_heading(doc, "3. 运行与部署环境", 1)
    add_heading(doc, "3.1 服务器端环境", 2)
    add_table(doc, ["类别", "环境"], [
        ("运行时", "Node.js 24.13.x"),
        ("Web 框架", "Fastify 5"),
        ("数据库", "PostgreSQL 17"),
        ("对象存储", "兼容 S3 API 的对象存储；生产环境采用 Cloudflare R2"),
        ("实时通信", "SSE、PostgreSQL LISTEN/NOTIFY、持久 journal/outbox"),
        ("可选支撑", "OpenTelemetry、Prometheus 指标、告警 Webhook、恶意文件扫描"),
    ], [2600, 7060])
    add_heading(doc, "3.2 浏览器端环境", 2)
    add_bullets(doc, [
        "支持 JavaScript、HTML5、CSS3 和 SSE 的现代桌面浏览器。",
        "支持现代移动浏览器；玩家端针对窄屏和触控导航进行适配。",
        "浏览器需允许访问对应站点并保持必要的会话 Cookie。",
    ])
    add_heading(doc, "3.3 生产入口", 2)
    add_table(doc, ["入口", "地址", "用途"], [
        ("官网", "https://getzhimu.com", "产品介绍、公开入口与内测申请"),
        ("创作者端", "https://app.getzhimu.com", "创作、内容管理、测试与账号设置"),
        ("主持端", "https://host.getzhimu.com", "运行房主持与现场控制"),
        ("玩家端", "https://play.getzhimu.com", "加入房间、阅读、调查与互动"),
        ("统一 API", "https://app.getzhimu.com/api", "身份、内容、运行态与实时事件"),
    ], [1600, 3600, 4460])

    add_heading(doc, "4. 系统组成与数据关系", 1)
    add_body(doc, "织幕由 Creator、Host、Player、官网和统一 API 五个表面组成。各前端保持角色专属工作流，通过同一套身份、权限和业务 API 访问 PostgreSQL 数据。Host 是唯一现场主持控制台，Creator 中不再维护第二套主持界面。")
    add_table(doc, ["组成", "主要职责", "核心数据"], [
        ("Creator", "创建和管理作品，设计机制，进行结构化创作、诊断、测试与发布准备", "世界、角色、章节、机制设计、线索、规则、资产"),
        ("Host", "选择运行房、查看提交汇总、推进阶段并执行权威结算", "房间、成员、机制运行态、提交、事件、检查点"),
        ("Player", "加入房间、查看个人内容、调查线索并提交机制选择", "角色绑定、阅读进度、线索、任务、自身提交状态"),
        ("官网", "展示产品、引导登录或内测申请", "公开内容和申请入口"),
        ("API", "认证、权限、业务写入、查询、实时同步和审计", "PostgreSQL、对象存储、事件日志"),
    ], [1600, 4100, 3960])
    add_body(doc, "世界模板的修改不会无条件覆盖已经运行中的房间。创建房间时，系统建立与作品版本对应的运行边界；主持操作和玩家行为记录为房间事件，便于同步、审计和复盘。")

    add_heading(doc, "5. 账号、权限与模型连接", 1)
    add_heading(doc, "5.1 注册与登录", 2)
    add_steps(doc, [
        "打开创作者端、主持端或玩家端入口。",
        "选择注册、登录或受支持的第三方登录方式；邮箱注册后输入邮件中的六位验证码，或使用邮件内的一键验证链接。",
        "六位验证码在十分钟内有效且只能使用一次；发送过于频繁或多次输错时，按页面提示等待或重新发送。",
        "登录成功后，系统建立安全会话并根据账号角色、世界成员关系和房间成员关系开放功能。",
        "如需退出，在账号菜单中执行退出；系统撤销当前会话。",
    ])
    add_body(doc, "权限判断由后端接口执行。前端隐藏按钮仅用于改善界面，不构成安全边界。账号可按内测状态、世界协作角色和房间角色获得不同访问范围。")
    add_heading(doc, "5.2 配置用户自备模型连接", 2)
    add_body(doc, "V1.0 的普通用户 AI 功能采用 BYOK 模式。用户在账号设置中配置自己的服务商、接口地址、模型名称与 API Key，调用消耗由用户在相应服务商的账户承担。平台自有 API Key 不进入用户模型池。")
    add_steps(doc, [
        "进入创作者端的“账号设置—模型连接”。",
        "选择 DeepSeek、OpenAI、OpenRouter、通义千问、智谱 GLM、SiliconFlow，或选择“OpenAI 兼容接口”。",
        "填写服务商要求的 API 地址、模型名称和 API Key；使用官方服务时优先保留系统建议的 API 地址。",
        "执行连接测试。测试成功后，将该连接设为当前使用连接。",
        "返回需要 AI 的创作或测试功能，从已启用模型中选择模型并执行任务。",
    ])
    add_notice(
        doc,
        "模型使用说明",
        "API Key 以 AES-256-GCM 加密保存，页面不会回显完整密钥。请不要在作品正文、反馈或截图中粘贴密钥。自定义兼容接口必须使用可信 HTTPS 地址；如余额不足、模型名错误或服务商限流，应在对应服务商处理。",
    )
    add_heading(doc, "5.3 三端用户资料、账号分级与内测权限", 2)
    add_body(doc, "同一账号可分别维护 Creator、Host、Player 三套显示名称和头像：创作者资料用于署名与协作，主持资料用于现场控制台和玩家可见的主持身份，玩家资料用于房间、广场、好友和消息。三端资料独立保存，修改一个入口不会覆盖另两个入口。")
    add_body(doc, "内测期间，系统以账号身份和已批准的内测申请为基础分配 beta 权益。管理员可审核申请并关联已注册账号；未获得对应权限的账号不能访问受限运营或创作能力。世界数量、单文件大小和存储容量由账号权益与配额约束；平台积分与商业套餐界面在本版本内保持关闭。")

    add_heading(doc, "6. 创作者端操作", 1)
    add_figure(doc, SCREENSHOTS["creator"], "创作者端的作品工作台与统一左侧导航", 1)
    add_heading(doc, "6.1 创建或选择世界", 2)
    add_steps(doc, [
        "登录创作者端，在世界列表中选择已有作品，或点击“新建世界”。",
        "填写作品名称、题材、人数、预计时长和简介；也可从模板开始。",
        "设置角色席位与章节基础结构。",
        "保存后进入创作驾驶舱，按当前阶段继续编辑。",
    ])
    add_heading(doc, "6.2 维护创作底座", 2)
    add_body(doc, "创作者可在创作宪法、真相链、案件时间线、伏笔和关系网络中维护作品的事实层。该信息用于检查后续章节、角色知识和线索编排是否自洽。")
    add_bullets(doc, [
        "创作宪法：记录题材、目标体验、禁区、核心冲突和创作原则。",
        "真相链与核诡：描述事件真相、因果链和关键反转。",
        "案件时间线：按时间记录真实事件和角色可见事件。",
        "伏笔与关系网：维护埋设位置、回收位置、人物关系及冲突。",
    ])
    add_heading(doc, "6.3 组装并确认活态故事总览", 2)
    add_body(doc, "创作驾驶舱的活态故事总览用于把分散的世界简介、灵感、创作宪法、角色、关系、核心事实和章节材料组合成统一故事脊柱。系统明确区分作者已确认、AI 暂拟和尚未解决内容；AI 结果先进入候选版本，不会直接覆盖作者设定。")
    add_steps(doc, [
        "在创作驾驶舱打开“当前故事总览”；首次使用时检查可参与装配的现有材料。",
        "已配置用户模型连接时，点击“组装第一版故事”或“重新组装候选”。",
        "核对一句话故事、整体故事、开场状态、引爆事件、核心冲突、玩家参与前提、行动循环、真相与转折。",
        "继续核对角色故事作用、章节因果主干、结局方向、待决定问题和暂定假设。",
        "比较候选版本与当前版本，选择采用或放弃；对确定内容逐项标记为“作者已确认”。",
        "后续角色创作、章节编排、机制设计和文稿生产以已采用的故事总览作为共同语义参考。",
    ])
    add_notice(
        doc,
        "作者控制边界",
        "重新装配候选版本时，系统保留已经标记为“作者已确认”的核心区块。候选版本只有在用户主动选择“采用为当前故事总览”后才写入作品设置。",
    )
    add_heading(doc, "6.4 设计并确认互动机制", 2)
    add_body(doc, "机制设计工作台用于在章节扩写和开局前明确玩家反复执行的动作、冲突来源、有限资源、即时反馈、失败推进、题材特异性与结局因果。作者选择线上表现后，系统同步生成玩家端提示、主持端操作提示和提交模式，避免三端各自解释同一机制。")
    add_steps(doc, [
        "在创作驾驶舱打开“机制设计工作台”，选择公开抉择、资源取舍、证据质证、顺序重建、限时危机或角色承诺。",
        "填写机制名称、一句话概述，并回答机制七问；描述世界内行为，不填写实现变量或内部状态键。",
        "检查玩家端说明与主持端说明的跨端预览；限时危机需在机制包中配置截止时间与超时默认方案。",
        "尚未定稿时点击“保存草稿”；需要进入后续生成上下文时点击“确认并用于生成”。",
        "退出存在未保存修改的工作台时，按二次确认放弃，避免误丢作者输入。",
    ])
    add_notice(
        doc,
        "机制作者边界",
        "草稿会明确标记为未确认，不得在生成过程中被补写成既定事实。玩家端只接收公开说明和临时句柄，不接收作者内部机制键、资源键或隐藏效果。",
    )
    add_heading(doc, "6.5 编辑角色、章节和私人内容", 2)
    add_steps(doc, [
        "进入角色管理，创建角色并填写公开档案、目标、关系和席位信息。",
        "进入章节编辑，建立章节顺序和段落内容。",
        "为每个角色维护私人剧本、秘密、任务和可见条件。",
        "进入独立玩家端，核对当前角色只看到被授权内容。",
        "保存修订并检查发布准备状态。",
    ])
    add_heading(doc, "6.6 场景、线索、物品与调查点", 2)
    add_body(doc, "场景与线索编排用于把正文内容转换成可执行的调查结构。线索可关联场景、角色、章节、获取条件和可见范围；物品可在主持发放、玩家持有和规则触发之间流转。")
    add_steps(doc, [
        "创建场景并设置名称、说明和所在阶段。",
        "创建线索或物品，填写公开文本、私有说明和归属。",
        "将线索关联到场景、调查点或结构化规则。",
        "在故事图谱中检查对象关系，避免孤立线索或无法触发的路径。",
        "使用发布检查核对缺失字段、可见性和引用完整性。",
    ])
    add_heading(doc, "6.7 结构化规则与小游戏", 2)
    add_body(doc, "规则由受控的触发条件和动作组成，例如在特定阶段、玩家完成调查或主持确认后发放线索、解锁章节或更新状态。系统对规则结构进行校验，不执行用户输入的任意脚本。")
    add_steps(doc, [
        "选择规则触发事件和适用对象。",
        "添加条件，例如阶段、角色、线索持有或主持确认状态。",
        "添加动作，例如发放线索、解锁内容、记录事件或启动小游戏。",
        "预览规则结果并修正结构错误。",
        "保存规则，并在测试房中验证触发与幂等行为。",
    ])
    add_heading(doc, "6.8 语义大纲与故事助手", 2)
    add_body(doc, "故事助手可按创作输入生成结构化大纲。V1.0 的语义质量门禁会检查玩家身份覆盖、章节因果与转折、题材专属资源、证据类型与来源、结局条件、角色责任和批次相似性；不满足结构合同的结果会进入校验或修复流程，而不是直接当作可发布内容。")
    add_steps(doc, [
        "在故事助手或创作向导中填写题材、玩家人数、核心冲突、限制条件和目标体验。",
        "选择已启用的用户模型连接并启动大纲生成。",
        "检查生成结果中的玩家身份、章节行动、状态与资源、证据、结局和角色责任。",
        "根据校验问题修订输入或候选大纲；只有通过人工核对的内容才继续进入角色脚本和章节生产。",
    ])
    add_heading(doc, "6.9 作品诊断与 AI 玩家测试", 2)
    add_body(doc, "作品诊断中心汇总结构完整性、角色信息差、伏笔回收、线索可达性和发布准备问题。机器压力测试在用户已配置的模型连接上运行，用隔离上下文检查理解、推理和行动中的结构性卡点，并把发现的问题反馈到创作修改流程。")
    add_steps(doc, [
        "先完成用户模型连接测试，并选中本次使用的模型。",
        "进入作品诊断中心，运行结构检查并查看按严重程度归类的问题。",
        "进入 AI 玩家测试，选择角色组合、测试范围和轮次。",
        "启动测试并查看各角色视角的理解、推理、卡点和潜在剧透。",
        "将问题转化为修改任务；修改后重新执行诊断或回归测试。",
    ])
    doc.add_page_break()
    add_heading(doc, "6.10 创建运行房与发布准备", 2)
    add_steps(doc, [
        "完成作品的必填内容和发布检查。",
        "选择“创建运行房”，填写房间名称、人数、可加入方式等信息。",
        "确认房间绑定的作品版本和初始章节。",
        "生成邀请码或邀请链接，并发送给玩家。",
        "转到主持端进行角色分配和开局前检查。",
    ])

    add_heading(doc, "7. 主持端操作", 1)
    add_figure(doc, SCREENSHOTS["host"], "主持端的运行房命令中心", 2)
    add_heading(doc, "7.1 进入运行房", 2)
    add_steps(doc, [
        "打开主持端并登录拥有该房间主持权限的账号。",
        "在房间列表选择目标运行房。",
        "检查成员数量、角色绑定、在线状态和当前阶段。",
        "确认所有必要内容已加载后进入命令中心。",
    ])
    add_heading(doc, "7.2 现场推进", 2)
    add_bullets(doc, [
        "查看玩家阅读进度、在线状态和异常提示。",
        "分配或调整角色，催促尚未完成准备的玩家。",
        "推进章节或阶段，并执行待确认事项。",
        "发放线索、物品或任务，解锁受条件控制的内容。",
        "启动或控制小游戏，预览并触发结构化规则。",
        "查看事件日志和运行报告；必要时创建检查点。",
    ])
    add_heading(doc, "7.3 查看提交并执行权威结算", 2)
    add_body(doc, "主持端的“机制运行”区域显示当前轮次、资源与证据、待决选择、玩家提交汇总和作者定义的截止时间。公开倾向可按选项查看数量与角色分布；角色承诺类提交属于秘密信息，仅主持端显示承诺人与内容。玩家提交只提供决策参考，不自动改变房间运行态。")
    add_steps(doc, [
        "核对当前轮次、作品绑定状态、剩余时间及每个选项的公开代价和风险。",
        "查看已收到的玩家倾向或秘密承诺；不要替未提交玩家补填选择。",
        "在截止前选择正常方案；限时机制到期后仅点击作者预设的“按默认方案结算”。",
        "提交时使用页面显示的当前运行修订；若收到修订冲突，刷新后重新核对，不用旧页面覆盖新结果。",
        "服务端校验主持权限、作品绑定、截止时间与期望修订号，写入唯一权威状态并记录主持审计。",
        "结算成功后检查资源、证据、结局候选和新轮次状态，再继续推进。",
    ])
    add_notice(
        doc,
        "限时结算约束",
        "服务器轮次时钟是截止时间真相源。系统阻止到期前提前执行默认后果，也阻止到期后继续执行普通选项，避免不同客户端本地时间造成双重结算。",
    )
    add_heading(doc, "7.4 异常处置", 2)
    add_body(doc, "如玩家断线，主持端保留其成员与进度状态；玩家重新连接后可恢复。对于重复点击或网络重试，关键写入使用幂等机制避免重复发放。主持人可踢出异常成员、撤销当前会话权限或恢复到已保存的检查点。")

    add_heading(doc, "8. 玩家端操作", 1)
    add_figure(doc, SCREENSHOTS["play"], "玩家端的角色主页、线索和互动入口", 3)
    add_heading(doc, "8.1 加入房间与选择角色", 2)
    add_steps(doc, [
        "打开玩家端，通过邀请链接或邀请码进入加入页面。",
        "登录或按房间允许的方式建立玩家身份。",
        "确认房间信息，选择可用角色或等待主持人分配。",
        "角色绑定后进入个人主页，阅读角色说明和当前章节。",
    ])
    add_heading(doc, "8.2 局中阅读与调查", 2)
    add_bullets(doc, [
        "在角色页查看个人档案、目标、关系和被授权的私密内容。",
        "在章节页阅读已解锁的角色剧本并保存阅读进度。",
        "在线索页查看已获得线索、来源、备注和可分享范围。",
        "在调查点提交调查，等待系统规则或主持人返回结果。",
        "在物品与任务页查看持有物、任务状态和完成条件。",
    ])
    add_heading(doc, "8.3 提交机制决策", 2)
    add_body(doc, "当前轮次存在机制决策时，玩家主页显示公开问题、选项说明、代价、风险、截止时间和自身提交状态。玩家只能提交服务端投影出来的临时决策与选项句柄，不能直接提交作者内部键或机制效果。")
    add_steps(doc, [
        "阅读当前问题和所有公开选项；限时机制先确认剩余时间。",
        "选择一个方案并提交。公开抉择、资源取舍等作为倾向提交；角色承诺按秘密模式提交。",
        "服务端校验房间成员身份、角色绑定、当前运行修订和截止时间后保存。",
        "截止前可再次提交，以同一玩家和角色的最新选择为准；页面会显示自身最新提交。",
        "提交后等待主持最终结算；玩家数量较多或多数倾向都不会绕过主持确认。",
    ])
    add_heading(doc, "8.4 投票、私密行动与社交", 2)
    add_body(doc, "玩家可按当前阶段参与投票、提交私密行动、更新怀疑度或口供，并在房间允许时使用广场、语音或其他互动入口。私密内容仅发送给有权查看的玩家或主持人。")
    add_heading(doc, "8.5 断线恢复", 2)
    add_body(doc, "玩家端通过实时事件游标跟踪房间变化。短暂断线后客户端重新连接并从游标继续接收；若游标失效，客户端重新读取当前房间状态，避免仅依赖内存中的旧数据。")
    add_steps(doc, [
        "网络恢复后先等待客户端自动重连，并观察当前章节、线索和成员状态是否刷新。",
        "如页面仍显示离线，手动刷新浏览器并重新进入原运行房。",
        "重新进入后核对角色绑定、阅读进度和最近一次提交结果，不要立即重复提交私密行动或投票。",
        "如状态仍不一致，停止继续操作并联系主持人，以服务器状态和房间事件日志为准。",
    ])
    add_notice(
        doc,
        "玩家端使用提示",
        "不要把邀请码、角色私密内容或模型 API Key 转发到公开页面。使用公共设备后应退出账号；共享截图前先遮盖个人信息和房间邀请码。",
    )

    add_heading(doc, "9. 三端实时同步", 1)
    add_body(doc, "Host 推进、玩家行动和规则触发等关键操作先写入 PostgreSQL，并记录可持久化事件。SSE 将事件推送到已授权客户端；PostgreSQL LISTEN/NOTIFY 用于多实例通知，轮询补偿用于处理通知丢失。")
    add_table(doc, ["事件示例", "主持端表现", "玩家端表现"], [
        ("章节推进", "当前阶段更新，记录操作结果", "解锁对应角色章节并刷新阅读页"),
        ("线索发放", "显示目标玩家和发放状态", "目标玩家收到新线索提示"),
        ("玩家提交行动", "进入待处理或事件日志", "显示已提交状态，避免重复提交"),
        ("玩家提交机制倾向", "提交汇总更新；秘密承诺仅主持可见", "仅显示自身最新提交，不改变最终运行态"),
        ("主持机制结算", "修订号递增并记录审计动作", "接收公开结果、资源变化或新轮次状态"),
        ("规则触发", "显示规则和动作结果", "按授权范围收到解锁或状态更新"),
        ("断线重连", "成员在线状态恢复", "从游标续接或重新读取当前状态"),
    ], [2200, 3730, 3730])

    add_heading(doc, "10. 存档、复盘与归档", 1)
    add_body(doc, "主持人可在关键节点创建检查点。检查点用于在允许范围内恢复房间状态；事件日志仍用于保留操作轨迹。房间结束后，系统根据事件、阶段、线索、投票和任务结果生成复盘材料。")
    add_figure(doc, SCREENSHOTS["archive"], "房间结束后的复盘与归档页面", 4)
    add_steps(doc, [
        "主持人确认房间进入结束状态。",
        "系统汇总关键事件、线索流转、角色行动和结果。",
        "生成复盘摘要或叙事，并由主持人检查内容。",
        "玩家在被授权范围内查看个人或公共复盘。",
        "运营或创作者可查看归档并将测试发现反馈到作品修订。",
    ])

    add_heading(doc, "11. 资产、导入与运营管理", 1)
    add_heading(doc, "11.1 资产与文档导入", 2)
    add_body(doc, "创作者可上传封面、图片和创作资料，或导入文本、PDF 和内容包以辅助结构化编辑。生产上传策略要求文件类型、单文件大小、账号剩余存储配额和安全检查；上传前预留容量，成功确认后转为实际占用，失败或取消时释放预留。对象存储中的文件按所有者和业务对象关联。")
    add_steps(doc, [
        "选择目标世界或资产目录。",
        "上传文件并等待类型、大小和安全检查。",
        "为文件填写用途和关联对象。",
        "如执行文档导入，检查系统提取的章节、角色或段落结构。",
        "确认后再写入作品；不正确的提取结果应放弃或重新映射。",
    ])
    add_heading(doc, "11.2 内测与运营管理", 2)
    add_bullets(doc, [
        "接收和审核内测申请，关联已注册用户并开放 beta 权益。",
        "按账号、世界成员和房间角色执行分级权限。",
        "管理 Creator、Host、Player 三端用户资料，并在账号导出或删除流程中同步处理资料头像。",
        "查看账号权益、世界数量和资产存储配额使用情况。",
        "查看用户反馈、内容审核状态和结构化审计记录。",
        "查看健康检查、就绪状态、指标、链路和告警配置。",
        "平台积分、商业套餐和平台自有用户 API 额度在 V1.0 内不启用。",
    ])

    add_heading(doc, "12. 推荐的完整操作流程", 1)
    add_steps(doc, [
        "注册创作者账号并完成内测权限开通。",
        "在账号设置中配置并测试用户自备模型连接；如不使用 AI，可跳过。",
        "创建世界，确定题材、角色数、章节和目标体验。",
        "维护创作宪法、真相链、时间线、伏笔和关系网。",
        "在机制设计工作台选择线上表现并完成机制七问；将定稿内容确认为生成上下文。",
        "编辑角色档案、私人剧本、章节和内容段落。",
        "编排场景、线索、物品、调查点、规则和小游戏。",
        "运行作品诊断并进入独立玩家端验收；需要时执行机器压力测试。",
        "修正阻断问题，通过发布检查后创建运行房。",
        "玩家加入并绑定角色，主持人在 Host 完成开局前检查。",
        "主持推进阶段，玩家阅读、调查并提交机制倾向；主持查看汇总后执行权威结算，系统实时同步公开结果。",
        "在关键节点创建检查点，出现异常时按权限恢复。",
        "结束房间并生成复盘，将真实运行发现反馈到下一次作品修订。",
    ])

    add_heading(doc, "13. 常见异常与处理", 1)
    add_table(doc, ["现象", "可能原因", "处理方法"], [
        ("无法登录", "凭据错误、邮箱未验证、会话失效或网络异常", "重试登录；完成验证；清理失效会话；检查网络"),
        ("无权访问世界或房间", "账号未加入、角色不符或内测权限未开通", "由所有者、主持人或管理员检查成员和账号分级"),
        ("模型连接测试失败", "Key 无效、余额不足、模型名或 API 地址错误、服务商限流", "到服务商控制台核对；修正配置后重新测试"),
        ("AI 功能不可用", "未启用用户连接，或当前功能不支持所选模型", "启用一个已测试连接并重新选择模型"),
        ("收不到实时更新", "网络断开、浏览器后台限制或 SSE 重连中", "等待自动重连；刷新页面并重新进入房间"),
        ("线索重复提交或发放", "网络重试或重复点击", "检查事件日志；系统幂等保护后以服务器结果为准"),
        ("机制提交被拒绝", "页面修订过旧、未绑定角色或限时选择已截止", "刷新当前房间状态，核对角色与剩余时间后再提交"),
        ("主持无法结算普通选项", "限时机制已到期或作品绑定已变化", "按作者预设默认方案结算；绑定变化时停止操作并执行恢复流程"),
        ("主持结算发生修订冲突", "其他主持操作已先更新权威状态", "重新读取运行态与提交汇总后再决定，不覆盖新结果"),
        ("上传失败", "格式、大小、安全扫描或对象存储异常", "更换合规文件；稍后重试；联系管理员查看扫描结果"),
        ("恢复后内容不一致", "检查点范围与当前事件边界不同", "停止继续操作，由主持人核对检查点和事件日志"),
    ], [2100, 3200, 4360])

    add_heading(doc, "14. 安全与隐私说明", 1)
    add_bullets(doc, [
        "身份和权限在后端校验，不依赖前端隐藏功能。",
        "浏览器会话使用安全 Cookie；敏感接口执行限流和审计。",
        "邮箱注册支持六位一次性验证码、有效期、重发冷却和错误尝试限制；验证失败不会清除已存在的其他有效会话。",
        "用户模型 API Key 使用 AES-256-GCM 加密保存，页面不回显完整值。",
        "平台自有模型密钥不进入普通用户调用池；系统级任务与用户任务保持路由边界。",
        "自定义模型接口执行地址校验和 SSRF 防护，用户应仅连接可信 HTTPS 服务。",
        "上传文件执行类型、大小和恶意内容检查；生产对象存储按业务权限访问。",
        "产品 HTML 写入收敛到受控安全 DOM 接口，CSP enforce 与 Trusted Types 门禁限制未授权脚本和不安全页面内容。",
        "实时事件按房间和受众过滤，私密行动与角色内容不会广播给无权用户。",
        "机制提交事件只公开决策提交数量；角色承诺的身份与内容仅主持端可见，玩家端不接收作者内部机制键。",
        "账号支持数据导出、会话撤销与删除流程；后台任务清理关联对象存储数据。",
    ])
    add_notice(
        doc,
        "使用者责任",
        "请勿上传没有授权的作品、个人敏感信息或违法内容。模型服务商会按其自身条款处理用户发出的内容；使用者应在调用前核对服务商的数据与隐私政策。",
        warning=True,
    )

    add_heading(doc, "15. 维护、部署与版本说明", 1)
    add_body(doc, "软件生产部署由 Creator/API、Host、Player 和官网分别发布。数据库变更通过顺序迁移执行；发布前进行构建、自动化测试、数据备份和关键流程验证。健康检查、就绪检查、指标、链路追踪、告警和结构化审计用于维护。")
    add_table(doc, ["项目", "V1.0 说明"], [
        ("版本号", VERSION),
        ("软件定位", "线上长线剧本杀与跑团的创作、主持和玩家协作平台"),
        ("用户侧 AI", "用户自备 API Key，可选择多家模型服务或 OpenAI 兼容接口"),
        ("故事语义层", "活态故事总览、作者确认状态、语义大纲生成与结构化质量校验"),
        ("机制三端闭环", "创作者机制设计、玩家倾向/秘密承诺提交、主持端权威结算与服务器截止时间"),
        ("账号资料", "Creator、Host、Player 三端显示名称与头像独立维护"),
        ("平台积分/付费套餐", "底座预留，当前不对普通用户开放"),
        ("平台自有 API 额度", "不进入用户模型池，仅可用于明确隔离的系统任务"),
        ("主要部署", "Creator/API：Railway；Host、Player、官网：Cloudflare Pages"),
        ("数据存储", "PostgreSQL 17 与兼容 S3 API 的对象存储"),
    ], [2800, 6860])
    add_body(doc, "本说明书用于软件著作权登记文档鉴别材料。软件界面在后续维护中可进行不改变核心功能的文字与视觉调整。")

    doc.save(path)
    return path


def parse_source_pages() -> list[list[str]]:
    text = SOURCE_TEXT.read_text(encoding="utf-8")
    lines = [line for line in text.splitlines() if re.match(r"^\d{2}:", line)]
    if len(lines) != 3000:
        raise RuntimeError(f"源程序交存稿应包含 3000 行，当前为 {len(lines)} 行")
    pages = [lines[idx:idx + 50] for idx in range(0, 3000, 50)]
    if len(pages) != 60 or any(len(page) != 50 for page in pages):
        raise RuntimeError("源程序分页失败")
    return pages


def build_source_code() -> Path:
    path = OUTPUT / "03_织幕源程序鉴别材料_V1.0.docx"
    pages = parse_source_pages()
    doc = Document()
    set_repeatable_metadata(
        doc,
        f"{SOFTWARE_NAME} {VERSION} 源程序鉴别材料",
        "普通交存：前 30 页和后 30 页，每页 50 行",
    )
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    section.header_distance = Inches(0.20)
    section.footer_distance = Inches(0.20)

    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    normal.font.size = Pt(7.5)
    normal.font.color.rgb = RGBColor.from_string("111827")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(f"{SOFTWARE_NAME}  {VERSION}  源程序鉴别材料")
    set_run_font(hr, size=7.5, bold=True, color=DARK_BLUE)
    hp.paragraph_format.space_after = Pt(0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("第 ")
    set_run_font(r, size=7.5, color=MUTED)
    page_run = add_field(fp, "PAGE", "1")
    set_run_font(page_run, size=7.5, color=MUTED)
    r = fp.add_run(" 页 / 共 60 页")
    set_run_font(r, size=7.5, color=MUTED)

    for page_index, page_lines in enumerate(pages):
        side = "前 30 页" if page_index < 30 else "后 30 页"
        label = doc.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label.paragraph_format.space_before = Pt(0)
        label.paragraph_format.space_after = Pt(2)
        label.paragraph_format.line_spacing = 1
        r = label.add_run(f"{side} · 第 {page_index + 1} 页")
        set_run_font(r, size=6.5, color=MUTED)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10)
        p.paragraph_format.keep_together = True
        for idx, source_line in enumerate(page_lines):
            r = p.add_run(source_line)
            set_run_font(r, name="Consolas", size=7.5, color="111827")
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if idx != 49:
                r.add_break()
        if page_index != 59:
            p.add_run().add_break(WD_BREAK.PAGE)

    doc.save(path)
    return path


def build_checklist() -> Path:
    path = OUTPUT / "04_织幕软著申请提交清单.docx"
    source_file_count, source_line_count = source_stats()
    doc = new_document(
        "织幕软件著作权申请提交清单",
        "提交步骤、材料一致性和待补充信息核对",
        "织幕 V1.0 · 提交清单",
    )
    add_title_page(
        doc,
        "软著申请提交清单",
        "材料核对 · 信息补充 · 递交顺序",
    )
    add_notice(
        doc,
        "这套材料已经完成什么",
        "已完成登记信息底稿、操作说明书、源程序前后各 30 页交存稿和本提交清单。相对 2026 年 8 月 6 日上一版，本轮补入机制设计工作台、玩家决策提交和主持端权威结算证据；当前建议以 2026 年 8 月 7 日为 V1.0 开发完成日期，并按网站实际公开情况选择发表状态。申请主体、首次发表日期、证件和真实权属仍须人工确认。",
    )

    add_heading(doc, "一、本材料包内容", 1)
    add_table(doc, ["序号", "文件", "用途", "状态"], [
        ("1", "01_织幕软件著作权登记信息表_V1.0.docx", "填写正式申请表时使用的资料底稿", "需补充身份和权属"),
        ("2", "02_织幕软件操作说明书_V1.0.docx / PDF", "文档鉴别材料", "已编制，提交前核对"),
        ("3", "03_织幕源程序鉴别材料_V1.0.docx / PDF", "程序鉴别材料，60 页 × 50 行", "已编制并脱敏"),
        ("4", "04_织幕软著申请提交清单.docx", "提交前逐项核对", "当前文件"),
    ], [700, 4000, 3300, 1660])

    add_heading(doc, "二、还需要你补充的材料", 1)
    add_table(doc, ["完成", "材料或信息", "要求"], [
        ("□", "著作权人名称和主体类型", "与身份证或营业执照完全一致"),
        ("□", "身份证明文件", "个人身份证；企业营业执照等有效证明"),
        ("□", "联系地址、邮编、联系人、手机和邮箱", "能够接收补正和登记通知"),
        ("□", "V1.0 开发完成日期", "当前建议 2026 年 8 月 7 日；按本轮机制功能完成提交和真实冻结情况确认"),
        ("□", "发表状态", "网站已公开提供软件功能时建议填“已发表”；补充首次日期，地点建议“中国／互联网”"),
        ("□", "权属情况", "独立/合作/委托/职务开发、原始/继受取得、权利范围"),
        ("□", "特殊权属证明", "如适用，提供合同、任务书、劳动关系、转让或授权文件"),
        ("□", "代理材料", "委托代理机构时按其要求提供授权或委托文件"),
    ], [850, 3700, 5110])

    add_heading(doc, "三、官方材料组成核对", 1)
    add_body(doc, "按现行软件著作权登记办法，通常需要软件著作权登记申请表、软件鉴别材料及相关证明文件。普通交存的程序和文档一般提交前、后各连续 30 页；不足 60 页的应提交全部。程序通常每页不少于 50 行，文档每页不少于 30 行。")
    add_table(doc, ["完成", "项目", "织幕 V1.0 对应材料"], [
        ("□", "正式登记申请表", "在官方系统填写；以 01 信息表为底稿"),
        ("□", "程序鉴别材料", "03 源程序鉴别材料，共 60 页，每页 50 行"),
        ("□", "文档鉴别材料", "02 操作说明书；如受理端要求截取，按其指引处理"),
        ("□", "申请人身份证明", "由申请人提供"),
        ("□", "权属证明", "仅在合作、委托、职务、受让等适用情形提供"),
        ("□", "中文译本", "外文证明文件适用"),
    ], [850, 2900, 5910])
    add_notice(
        doc,
        "规则来源",
        "《计算机软件著作权登记办法》要求使用统一表格、中文填写，并规定程序和文档鉴别材料的常规交存方式。不同地区、代理机构或受理系统的电子文件格式可能有操作差异，上传时以受理页面的最新提示为准。",
    )

    add_heading(doc, "四、一致性检查", 1)
    add_table(doc, ["完成", "检查项", "统一值"], [
        ("□", "软件全称", SOFTWARE_NAME),
        ("□", "软件简称", SHORT_NAME),
        ("□", "版本号", VERSION),
        ("□", "开发完成日期", "建议 2026 年 8 月 7 日；确认后在全部申请信息中保持一致"),
        ("□", "著作权人名称", "【待填写后，与证件完全一致】"),
        ("□", "申请表功能说明与说明书", "仅描述 V1.0 当前功能，不把未来规划写成已上线"),
        ("□", "源程序页眉与说明书封面", f"{SOFTWARE_NAME} {VERSION}"),
        ("□", "发表状态", "建议已发表；首次发表日期待确认，并在全部申请信息中保持一致"),
    ], [850, 3150, 5660])

    add_heading(doc, "五、相对上一版新增功能证据", 1)
    add_body(doc, f"对照基线为 {PREVIOUS_PACKAGE}。本轮功能实现提交完成于 {FEATURE_IMPLEMENTED_DATE_CN}，材料于 {PACKAGE_DATE_CN} 重新生成。下表同时给出功能闭环、关键代码和源程序交存稿可见页；玩家提交始终是倾向或秘密承诺，最终状态由主持端权威结算。")
    add_table(doc, ["新增功能", "可验证行为", "关键代码与交存页"], [
        (
            "机制设计工作台",
            "选择六类线上表现；回答机制七问；保存草稿或确认为生成上下文；预览玩家/主持说明。",
            f"creator workbench、mechanism design/interactions；{evidence_page_label(FEATURE_EVIDENCE['机制设计工作台'])}",
        ),
        (
            "玩家决策提交",
            "校验成员、角色、运行修订与截止时间；公开倾向/秘密承诺可在截止前以最新提交覆盖。",
            f"player controller/view、submission service/routes、迁移 107/108；{evidence_page_label(FEATURE_EVIDENCE['玩家决策提交'])}",
        ),
        (
            "主持端权威结算",
            "查看提交汇总和服务器时钟；期望修订校验；到期前禁默认、到期后仅允许作者预设默认方案；记录审计与状态事件。",
            f"host workspace/controller、runtime service/routes；{evidence_page_label(FEATURE_EVIDENCE['主持端权威结算'])}",
        ),
        (
            "跨端同步与隐私",
            "提交数量与公开结果按房间事件同步；秘密承诺详情仅主持可见，玩家投影隐藏作者内部机制键。",
            "room event audience、mechanism runtime projection 与跨端测试证据",
        ),
    ], [2100, 4100, 3460])

    add_heading(doc, "六、源程序材料检查", 1)
    add_bullets(doc, [
        "本次采用普通交存，程序为前 30 页和后 30 页，共 60 页，每页 50 行。",
        "纳入 Creator、Host、Player、共享模块、后端业务代码和数据库迁移。",
        f"本次整理脚本统计 {source_file_count} 个自研源文件、约 {source_line_count:,} 行源代码；交存稿截取前后各 1,500 行。",
        "不纳入 node_modules、第三方依赖、构建产物、测试文件、环境文件、日志和临时文件。",
        "交存稿已对疑似 token、secret、password、api_key 和数据库连接串进行脱敏。",
        "页面按 A4、等宽字体、连续页码输出；提交 PDF 前检查没有换行导致的额外页。",
        "前 30 页与后 30 页分别从确定性文件序列连续截取；相邻页应与源文本相邻 50 行一致，不跳行、不重复。",
        "前段包含机制设计工作台与共享交互合同，后段包含玩家提交、主持端操作及服务端权威结算。",
        "不要把 .env、部署平台密钥、模型 API Key、Cookie 签名或私钥追加进材料。",
    ])

    add_heading(doc, "七、文档材料检查", 1)
    add_bullets(doc, [
        "说明书封面、页眉和版本信息与申请表一致。",
        "说明书只写实际已实现功能；机制设计、玩家决策提交、主持权威结算、活态故事总览、语义大纲、三端用户资料、邮箱验证码和资产配额与当前代码一致。",
        "截图不含真实用户隐私、API Key、数据库连接串或内部运维凭据。",
        "操作步骤覆盖创作者端、主持端、玩家端、模型连接、实时同步、复盘与异常处理。",
        "如上传系统要求文档前后各 30 页，则在最终 PDF 页面数确定后按其提示截取；不足 60 页提交全部。",
    ])

    add_heading(doc, "八、推荐递交顺序", 1)
    add_steps(doc, [
        "先在 01 信息表补齐申请主体、完成日期、发表状态和权属。",
        "核对软件名称与版本后，阅读 02 说明书并删改任何与你真实情况不符的描述。",
        "打开 03 源程序 PDF，确认恰好 60 页，页码连续，首末页内容可读。",
        "准备身份证明及适用的合同、任务书、授权或转让证明。",
        "登录中国版权保护中心相关登记系统，按 01 信息表填写正式申请表。",
        "上传申请表、源程序、说明书及证明文件，并按系统提示完成签章或确认。",
        "保存提交回执和受理编号；收到补正通知时，只修改被要求的部分并再次做一致性检查。",
    ])

    add_heading(doc, "九、最终签发页", 1)
    add_table(doc, ["核对项目", "签发"], [
        ("材料信息已补齐", "□ 是"),
        ("身份与权属证明已准备", "□ 是"),
        ("申请表、源程序、说明书名称和版本一致", "□ 是"),
        ("源程序 PDF 为 60 页且无密钥", "□ 是"),
        ("说明书内容与实际 V1.0 一致", "□ 是"),
        ("申请人/经办人", "____________________________"),
        ("签发日期", "________ 年 ____ 月 ____ 日"),
    ], [4200, 5460])
    add_body(doc, "说明：本清单用于材料整理，不构成法律意见。对于合作、委托、职务开发或权利转让等复杂权属情况，建议在提交前由专业人员核对证明链。")

    doc.save(path)
    return path


def build_feature_evidence_report() -> Path:
    path = OUTPUT / "05_织幕V1.0_相对上一版新增功能证据.txt"
    page_map = source_evidence_pages()
    lines = [
        "织幕 V1.0 相对上一版新增功能证据",
        "================================",
        "",
        f"上一版基线：{PREVIOUS_PACKAGE}",
        f"功能实现提交完成：{FEATURE_IMPLEMENTED_DATE_CN}",
        f"本次材料编制：{PACKAGE_DATE_CN}",
        "",
        "一、新增功能闭环",
        "",
        "1. 机制设计工作台",
        "   - 创作者选择公开抉择、资源取舍、证据质证、顺序重建、限时危机或角色承诺。",
        "   - 回答机制七问，保存草稿或确认为生成上下文，并预览玩家端与主持端说明。",
        "2. 玩家决策提交",
        "   - 玩家提交公开倾向或秘密承诺；服务端校验成员、角色、运行修订与截止时间。",
        "   - 截止前重复提交以同一玩家/角色的最新选择覆盖；提交本身不改变权威运行态。",
        "3. 主持端权威结算",
        "   - 主持端查看提交汇总、服务器轮次时钟、资源、证据和作者预设默认方案。",
        "   - 服务端校验主持权限、作品绑定和期望修订号；到期前禁止提前默认，到期后仅允许默认方案。",
        "4. 跨端同步与隐私",
        "   - 提交数量和公开结果通过房间事件同步；秘密承诺详情仅主持端可见。",
        "   - 玩家投影隐藏作者内部机制键，只暴露临时句柄、公开说明和自身提交状态。",
        "",
        "二、源程序交存稿证据页",
        "",
    ]
    for feature, paths in FEATURE_EVIDENCE.items():
        lines.append(f"【{feature}】")
        for rel in paths:
            pages = page_map.get(rel, [])
            label = "、".join(f"第 {page} 页" for page in pages) if pages else "未落入本次前后 30 页摘录（保留为仓库证据）"
            lines.append(f"- {rel}：{label}")
        lines.append("")
    lines.extend([
        "连续性说明：前 30 页和后 30 页分别从确定性文件序列连续截取，每页 50 行，页际不跳行；证据文件未重复计入总文件数。",
        "",
        "三、自动化测试证据",
        "",
        *[f"- {rel}" for rel in FEATURE_TEST_EVIDENCE],
        "",
        "四、功能边界说明",
        "",
        "- 玩家提交是咨询性倾向或秘密承诺，不等同于多数表决自动生效。",
        "- 只有主持权限通过且期望修订、作品绑定、截止时间均有效的服务端操作才能写入权威机制运行态。",
        "- 本证据清单用于解释相对上一版的新增范围；正式交存仍以 03 源程序鉴别材料的连续 60 页为准。",
        "",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for old in OUTPUT.glob("*.docx"):
        old.unlink()
    created = [
        build_application_info(),
        build_manual(),
        build_source_code(),
        build_checklist(),
    ]
    evidence_report = build_feature_evidence_report()
    manifest = {
        "software": SOFTWARE_NAME,
        "version": VERSION,
        "generated": PACKAGE_DATE,
        "files": [str(path) for path in [*created, evidence_report]],
        "design_preset": "compact_reference_guide",
        "layout_override": "official_submission_a4",
        "source_pages": 60,
        "source_lines_per_page": 50,
        "source_files": source_stats()[0],
        "source_total_lines": source_stats()[1],
        "previous_package": PREVIOUS_PACKAGE,
        "feature_implemented_date": "2026-08-07",
        "feature_baseline": "story-spine, semantic-outline, portal-profiles, email-code-verification, asset-quota, mechanism-design-workbench, player-mechanism-submissions, host-authoritative-settlement",
        "feature_evidence": FEATURE_EVIDENCE,
        "feature_test_evidence": FEATURE_TEST_EVIDENCE,
        "source_evidence_pages": source_evidence_pages(),
        "source_continuity": "front and back excerpts are independently contiguous deterministic sequences, 50 lines per page",
    }
    (OUTPUT / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    readme = f"""织幕 V1.0 软件著作权申请材料（更新版）
========================================

材料编制日期：{PACKAGE_DATE_CN}

一、使用顺序

1. 先打开“01_织幕软件著作权登记信息表_V1.0.docx”。
2. 补齐全部【待填写】和【待确认】项目，包括申请主体、证件与联系方式、首次发表日期及真实权属；当前开发完成日期建议为 2026 年 8 月 7 日，发表状态建议按真实公开情况确认。
3. 阅读“02_织幕软件操作说明书_V1.0.docx”，确认机制设计、玩家决策提交、主持权威结算、活态故事总览、三端资料及资产配额等描述与实际情况一致。
4. “03_织幕源程序鉴别材料_V1.0”按前 30 页、后 30 页整理，共 60 页，每页 50 行。
5. 按“04_织幕软著申请提交清单.docx”准备身份证明和适用的权属证明。
6. 查看“05_织幕V1.0_相对上一版新增功能证据.txt”，核对本轮新增功能与交存页映射。

二、统一名称

软件全称：{SOFTWARE_NAME}
软件简称：{SHORT_NAME}
版本号：{VERSION}

三、当前版本边界

- 普通用户 AI 功能采用用户自备 API Key（BYOK）模式。
- 活态故事总览和语义大纲结果均需作者核对，AI 候选不会自动覆盖作者确认内容。
- 机制设计工作台保存作者机制七问与跨端交互合同；玩家提交仅表达倾向或秘密承诺，主持端经服务端校验后执行唯一权威结算。
- 限时机制以服务器轮次时钟为准，到期前不能提前执行默认后果，到期后仅按作者预设默认方案结算。
- 玩家端不接收作者内部机制键，秘密承诺详情仅主持端可见。
- Creator、Host、Player 可分别维护显示名称与头像。
- 平台积分、付费套餐和平台自有 API 额度当前不对普通用户开放。
- 本次源代码统计为 {source_stats()[0]} 个自研文件、约 {source_stats()[1]:,} 行；正式交存稿为 60 页 × 50 行。

四、提交前仍需人工完成

- 将申请主体、首次发表日期和权属填写到正式申请表；确认建议的 2026 年 8 月 7 日完成日期与发表状态，并保持全部材料一致。
- 核对 PDF 页码、截图隐私和源程序敏感信息扫描结果。
- 以中国版权保护中心受理页面的最新格式和签章要求为准。

说明：本材料包不替代官方在线申请表，也不构成法律意见。
"""
    (OUTPUT / "README_先看这里.txt").write_text(readme, encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
