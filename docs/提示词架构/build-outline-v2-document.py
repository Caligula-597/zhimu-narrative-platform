from __future__ import annotations

import json
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
OUT_FILE = OUT_DIR / "织幕剧情大纲提示词与校验框架_V2.2.docx"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
FONT_CODE = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "667085"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
GOLD = "9A6700"
RED = "9B1C1C"
GREEN = "216E39"
WHITE = "FFFFFF"
BLACK = "000000"
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name=FONT_LATIN, east_asia=FONT_CJK, size=None, bold=None,
                 italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill=LIGHT_FILL, border_color=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def make_numbering(doc, num_format, text_pattern, font_name=FONT_LATIN):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_format)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_pattern if num_format == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
    # Word requires every abstractNum to appear before the concrete num nodes.
    # Appending it after existing num nodes makes Word silently fall back to a
    # decimal list, which is especially visible in long Chinese bullet lists.
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    cleanup_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:numIdMacAtCleanup")
        ),
        len(numbering),
    )
    numbering.insert(cleanup_index, num)
    return num_id


def apply_num(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_bullet(doc, text, num_id, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, num_id, level)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=BLACK)
    return p


def add_numbered(doc, text, num_id, level=0):
    p = doc.add_paragraph()
    apply_num(p, num_id, level)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(text), color=BLACK)
    return p


def add_body(doc, text, bold_prefix=None, color=BLACK, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=color)
    else:
        set_run_font(p.add_run(text), color=color)
    return p


def add_callout(doc, label, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.18
    set_paragraph_shading(p, LIGHT_FILL, color)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    return p


def add_code_block(doc, code, label=None):
    if label:
        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(4)
        p_label.paragraph_format.space_after = Pt(3)
        set_run_font(p_label.add_run(label), size=9, bold=True, color=MUTED)
    for index, line in enumerate(code.splitlines() or [""]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F7F8FA")
        p_pr.append(shd)
        if index == 0:
            borders = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "4")
            top.set(qn("w:color"), "D0D5DD")
            borders.append(top)
            p_pr.append(borders)
        run = p.add_run(line if line else " ")
        set_run_font(run, name=FONT_CODE, east_asia=FONT_CJK, size=8.7, color="273444")
    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_after = Pt(8)
    p_end.paragraph_format.line_spacing = 1.0
    p_pr = p_end._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F8FA")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "D0D5DD")
    borders.append(bottom)
    p_pr.append(borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        set_run_font(p.add_run(header), size=9.3, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            set_run_font(p.add_run(str(value)), size=9.2, color=BLACK)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def extract_prompt():
    node = os.environ.get("CODEX_NODE_EXECUTABLE")
    if not node:
        raise RuntimeError("CODEX_NODE_EXECUTABLE is required")
    script = r'''
import { buildStoryOutlineMessages } from "./backend/src/prompts/outline.js";
const spec = {
  title: "{{title}}",
  playerCount: 6,
  chapterCount: 5,
  chapterKeys: ["chapter-1", "chapter-2", "chapter-3", "chapter-4", "chapter-5"],
  targetWordCount: "{{targetWordCount}}",
  wordsPerSectionMin: "{{wordsPerSectionMin}}",
  sceneCount: "{{sceneCount}}",
  investigationPointCount: "{{investigationPointCount}}",
  clueCount: "{{clueCount}}",
  constraints: ["{{constraint1}}", "{{constraint2}}"],
  notes: ["{{note1}}"]
};
const brief = {
  title: "{{title}}",
  premise: "{{premise}}",
  conflicts: "{{conflicts}}",
  style: "{{style}}",
  audience: "{{audience}}",
  roleRequirements: "{{roleRequirements}}",
  evaluationFocus: "{{evaluationFocus}}",
  playerCount: 6,
  chapterCount: 5
};
process.stdout.write(JSON.stringify(buildStoryOutlineMessages(brief, spec)));
'''
    result = subprocess.run(
        [node, "--input-type=module", "-e", script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    messages = json.loads(result.stdout)
    return messages[0]["content"], messages[1]["content"]


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    title = styles["Title"]
    title.font.name = FONT_LATIN
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(8)

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (11.5, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("织幕 · 剧情生成架构参考"), size=8.5, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("V2.2  ·  "), size=9, color=MUTED)
        add_page_field(p)


def build_document():
    system_prompt, user_prompt = extract_prompt()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    doc.core_properties.title = "织幕剧情大纲提示词与校验框架 V2.2"
    doc.core_properties.subject = "剧情大纲协议、题材贡献、实体资源、分支因果、结局路径与批次校准"
    doc.core_properties.author = "织幕"
    doc.core_properties.keywords = "织幕, 剧情大纲, 提示词, DeepSeek, 校验器, V2.2"

    bullet_num = make_numbering(doc, "bullet", "•", FONT_LATIN)

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(76)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("织幕 · 创作系统内部参考"), size=11, bold=True, color=GOLD)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("剧情大纲提示词与校验框架"), size=28, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run_font(subtitle.add_run("Player-driven Outline Protocol V2.2"), size=15, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    set_run_font(meta.add_run("版本：V2.2  ·  实现快照：2026-07-29"), size=10.5, bold=True, color=MUTED)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(56)
    set_run_font(meta2.add_run("覆盖生成提示词、作用性门禁、状态因果、修复分流与四十篇语义去重"), size=10, color=MUTED)
    add_callout(
        doc,
        "文档用途",
        "用于产品、创作与研发共同审阅。完整提示词从当前源码动态提取；文档中的双花括号表示运行时变量，不包含任何 API 密钥或用户隐私。",
        BLUE,
    )
    doc.add_page_break()

    add_heading(doc, "1. 文档定位与结论", 1)
    add_body(
        doc,
        "本框架把大纲生成从“字段齐全的空洞 JSON”升级为“贡献按题材适配、实体资源可登记、条件失败可执行、玩家行动能沿因果路径抵达结局”的作用性协议。"
    )
    add_callout(
        doc,
        "核心原则",
        "提示词负责要求模型交付完整设计；校验器不只检查“有没有”，还检查字段是否被下游真正读取。自然语言 nextState、改名后的同源证据和机械填充的角色行动都不能代替真实因果。",
        GREEN,
    )
    for text_value in [
        "旧版 V1 大纲仍可读取，但 readiness.readyForExpansion 固定为 false。",
        "新版 V2.2 只有通过结构完整、作用性、可达性和模板门禁后，才能进入角色矩阵和正文扩写。",
        "角色不再统一绑定核心证据；贡献可落在关系、承诺、权限、资源、任务或风险，但必须形成通往结局的因果路径。",
        "DeepSeek 失败分成 patch 与 rebuild：patch 使用 0.25，rebuild 使用 0.55 且不携带完整旧 JSON，最多三次。",
        "十一维相似度当前是字符 bigram 的等权均值，默认只进入人工复核；待 100–200 对标注样本校准后再启用拒绝。",
    ]:
        add_bullet(doc, text_value, bullet_num)

    add_heading(doc, "2. 整体生成架构", 1)
    add_body(doc, "当前总纲链路分为八个明确阶段，每一阶段都有可机器判断的输入、输出和失败处理。")
    pipeline_rows = [
        ("1", "创作 Brief", "题目、梗概、文风、冲突、人数与章节数", "作为不可替换的创意源"),
        ("2", "规格归一化", "玩家数、章节 key、字数和约束", "生成 spec"),
        ("3", "V2.2 总纲生成", "系统提示词 + 不可信用户素材", "结构化 JSON"),
        ("4", "单篇作用性校验", "题材贡献、实体资源、分支传递、结局可达", "通过或返回失败项"),
        ("5", "修复模式分流", "patch 保留合格结构；rebuild 重做坏地基", "最多三次"),
        ("6", "可扩写标记", "全部门禁通过", "readyForExpansion=true"),
        ("7", "批次相似度复核", "十一类指纹、字符相似度与历史库", "问题、警告与 token 报告"),
        ("8", "后续生产", "结构、角色矩阵、私人分幕与正文", "进入下一流水线"),
    ]
    add_table(doc, ["阶段", "名称", "关键内容", "结果"], pipeline_rows, [600, 1550, 4300, 2910])

    add_heading(doc, "2.1 数据流", 2)
    flow = (
        "brief → normalizeStoryBrief → validateStorySpec → buildStoryOutlineMessages\n"
        "→ DeepSeek JSON → validateStoryOutline(V2.2, strict)\n"
        "→ [失败：repairMode=patch 定点修复 / rebuild 整体重构，最多 3 次]\n"
        "→ readiness.readyForExpansion=true\n"
        "→ validateOutlineBatchDiversity → 角色矩阵 / 公共结构 / 私人分幕"
    )
    add_code_block(doc, flow)

    add_heading(doc, "2.2 运行参数", 2)
    add_table(
        doc,
        ["参数", "当前值", "说明"],
        [
            ("模型", "deepseek-v4-flash", "由环境变量 DEEPSEEK_MODEL 覆盖"),
            ("单次最大输出", "12,000 tokens", "容纳完整玩家矩阵和证据图"),
            ("首次 temperature", "0.45", "保持创意差异"),
            ("patch temperature", "0.25", "稳定补字段、key 和登记错误"),
            ("rebuild temperature", "0.55", "重做坏地基，避免低温复刻旧结构"),
            ("最大尝试次数", "3", "DEEPSEEK_OUTLINE_MAX_ATTEMPTS，范围 1–3"),
            ("严格协议版本", "outlineVersion=2 + outlineRevision=2.2", "旧 V2.0/V2.1 只兼容读取"),
        ],
        [2100, 1850, 5410],
    )

    add_heading(doc, "2.3 准入结果", 2)
    add_table(
        doc,
        ["结果", "服务端状态", "后续动作"],
        [
            ("严格通过", "readyForExpansion=true", "允许进入角色矩阵、公共结构和私人分幕"),
            ("可读取旧版", "readyForExpansion=false", "仅展示或人工迁移，不进入自动扩写"),
            ("V2.2 校验失败", "DEEPSEEK_OUTPUT_INVALID + repairMode + issues", "按 patch/rebuild 分流"),
            ("三次仍失败", "向上返回最后一次错误", "不输出半成品，不生成批次成功汇总"),
        ],
        [1900, 2800, 4660],
    )

    add_heading(doc, "3. 完整生成提示词", 1)
    add_callout(
        doc,
        "使用说明",
        "以下内容由 backend/src/prompts/outline.js 在生成文档时动态展开。PRODUCT_BOUNDARY 已合并；双花括号为说明性占位符。用户素材通过 untrustedUserPayload 序列化，不能覆盖系统提示词。",
        BLUE,
    )
    add_heading(doc, "3.1 System Prompt（完整展开）", 2)
    add_code_block(doc, system_prompt, "当前源码导出")
    add_heading(doc, "3.2 User Prompt（模板化展开）", 2)
    add_code_block(doc, user_prompt, "运行时会把 spec 与 brief 替换为真实 JSON")

    add_heading(doc, "4. 大纲协议 V2.2 字段框架", 1)
    add_body(
        doc,
        "系统提示词已经包含完整 JSON schema。本节用于产品和创作人员快速理解每个结构为什么存在，以及它对应哪类旧问题。"
    )
    field_rows = [
        ("outlineVersion / outlineRevision", "协议版本", "必须为 2 / 2.2；旧 V2.0/V2.1 不自动冒充新门禁"),
        ("sourceFidelity", "原始创意忠实度", "标题逐字一致；至少两个梗概原文锚点绑定章节和支持项"),
        ("logline", "一句话冲突", "高概念入口，不允许结局降级"),
        ("truthTimeline", "唯一幕后真相", "完整因果顺序，仅创作者/主持人可见"),
        ("hookPromises", "梗概兑现清单", "每个异常承诺有 payoff 和至少两个题材适配支持项"),
        ("genreMechanic", "题材玩法机制", "规则、操作、限制、使用章节和最终兑现"),
        ("genreProfile", "题材门禁配置", "按推理、情感、政治、综艺、生存或混合选择章节进展方式"),
        ("entities", "实体注册表", "统一 NPC、机构、系统、设备、物证、地点和别名 key"),
        ("resources", "资源注册表", "声明初值、上下限、所有者、恢复性与数值类型"),
        ("players", "玩家角色矩阵", "人数精确；贡献锚点按题材选择并形成结局因果路径"),
        ("centralResponsibilityRoleKeys", "核心责任角色", "必须来自玩家，不允许 NPC 包办核心阴谋"),
        ("evidenceGraph", "题材可选证据图", "推理题材执行双源；非推理题材可为空"),
        ("misdirections", "题材适配误导", "嫌疑、记忆、阵营、公众叙事或风险误判"),
        ("chapterBeats", "章节行动", "状态读写、失败 variant、证据开关、资源变化和聚光节奏"),
        ("endingLogic", "累计结局", "可读取 state/resource/evidence，并验证可达性与裁决"),
        ("batchFingerprint", "批次剧情指纹", "十一维字符相似度报告；校准前默认人工复核"),
        ("suggestions", "可选扩写建议", "只能锦上添花，不得推迟必填设计"),
        ("readiness", "服务端追加状态", "表明是否通过严格门禁、能否进入扩写"),
    ]
    add_table(doc, ["字段", "用途", "硬性意义"], field_rows, [1950, 2200, 5210])

    add_heading(doc, "4.1 玩家角色的最低完整度", 2)
    for value in [
        "key：唯一、稳定，并被章节、证据和核心责任引用。",
        "name：具体姓名；禁止“角色A、玩家B、嘉宾1、队员C”等占位名。",
        "identity：说明专业能力或关系位置为什么不可替代。",
        "publicGoal / hiddenGoal：公开目标和隐藏目标必须形成真实冲突。",
        "coreSecret：必须与主线有因果关系，不能只是无关污点。",
        "exclusiveAnchorKey：必须指向 contribution.anchorType 对应的已登记 key。",
        "activePlan：必须写具体对象、方法和代价，不能只写调查、质问或交换信息。",
        "arc：角色从起点经选择到终点的变化。",
        "spotlightChapterKey：必须与 turnChapterKeys、chapterActions 和 triggerRoleKeys 对齐。",
        "contribution：按题材选择 evidence/relationship/commitment/authority/resource/task/risk 锚点。",
        "chapterActions：至少覆盖 ceil(章节数×0.6)；每项包含稳定 actionTargetKey、资源/状态/证据影响和他人影响。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "4.2 章节的最低完整度", 2)
    add_body(doc, "每章不是“发一份资料”，也不是一句“影响后续局势”。完整章节需要以结构化字段证明局面真的发生变化：")
    for value in [
        "triggerRoleKeys：哪些玩家发动本章关键事件。",
        "playerAction：可以在桌面或线上实际执行的行为。",
        "actionObject：明确被修改的系统、开启的空间、质证的材料、保护的人或消耗的资源。",
        "irreversibleConsequence：不能无代价复原的结果。",
        "stateReads / stateWrites：上一章写入的状态如何在本章读取，本章又怎样改变后续。",
        "entryConditionMode / onReadPass / onReadFail：条件成功和失败分别进入哪个 variant，失败支付什么代价。",
        "unlocksEvidenceKeys / locksEvidenceKeys：哪些证据因玩家行为被打开或永久关闭。",
        "resourceDeltas：只引用 resources 注册表，amount 使用数字，并遵守初值和上下限。",
        "progressMode：按题材选择证据、关系、承诺、资源、任务、观众或风险推进。",
        "decision：按 genreProfile 决定频率；有决策时至少两个选项并写入真实状态。",
        "nextState：只作为可读摘要，不能代替以上结构化因果。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5. 单篇硬校验器", 1)
    add_body(
        doc,
        "校验器不是评分器，而是准入门。只要存在一项失败，系统就不会返回 readyForExpansion=true。错误以 DEEPSEEK_OUTPUT_INVALID 和结构化 issues 返回给自动修复流程。"
    )

    add_heading(doc, "5.1 梗概忠实度与高概念兑现", 2)
    for value in [
        "sourceFidelity.briefTitle 必须逐字等于 brief.title。",
        "premiseElements 至少两项；element 必须原样出现在 brief.premise。",
        "每个 premiseElement 必须写明 implementation，并引用至少一个合法章节和 supportKey。",
        "hookPromises 至少一项；每项 promise 和 payoff 不得为空。",
        "每项 hookPromise 至少引用两个 supportKeys；mystery 才强制两条独立证据。",
        "payoff 过短会拒绝，避免用一句“其实是骗局”降级高概念。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.2 六人/多人角色结构", 2)
    for value in [
        "players 数量必须精确等于 spec.playerCount。",
        "key 和姓名都必须唯一；姓名不得使用占位模板。",
        "身份、目标、秘密、具体主动计划、人物弧光和聚光章均为必填。",
        "独占锚点必须与 contribution.anchorType 对应，可为证据、状态、资源或任务实体。",
        "每名玩家至少覆盖 ceil(章节数×0.6) 个关键章节，避免六人五章被强行填成 30 条空动作。",
        "每人至少拥有一次主线转折、一次影响其他玩家的行动，并通过状态、资源或证据开关形成结局因果路径。",
        "chapterActions 声明的 stateWriteKeys、resourceKeys 和 evidenceEffectKeys 必须在同章公共因果中真实发生。",
        "五章至少覆盖四个不同聚光章；每章原则上最多两名，三名时必须解释 sharedSpotlightConflict。",
        "centralResponsibilityRoleKeys 至少一项，且全部属于玩家 key。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.3 公平推理与证据来源", 2)
    for value in [
        "mystery 的证据 key 和核心结论 key 必须存在且唯一；非推理题材可以不建立证据图。",
        "provenanceGroup 必须引用 entities 中登记的稳定来源实体；originActorKey 引用玩家或实体。",
        "派生截图、转录或摘要必须通过 derivedFromEvidenceKeys 指向原证据；禁止自引用和派生循环。",
        "obtainedBy 必须描述玩家怎样取得证据，不能仅写“主持人发放”。",
        "每个核心结论至少引用两条证据。",
        "同一结论不仅要求 sourceType 不同，还要求 provenanceGroup 或派生根不同。",
        "若全部证据来自同一 originActorKey，人物口供和同一人物日记不能冒充双源。",
        "最终章若有单条证据一次性支持全部核心结论，直接拒绝。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.4 题材适配误导", 2)
    for value in [
        "mystery 至少两条 misdirections；其他题材至少一条。",
        "kind 随题材变化：嫌疑/证据、记忆/关系、阵营/权限、公众叙事/任务、风险/资源。",
        "每条必须包含 apparentInterpretation、trueCause、mainlineImpact 和 lastingConsequence。",
        "supportKeys 与 disproofKeys 可引用证据、状态、资源或实体；mystery 必须含具体证据。",
        "无关的赌债、婚外情、欠款或偷窃不能单独充当主线误导。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.5 玩家行动与状态因果", 2)
    for value in [
        "chapterBeats 必须精确覆盖全部 spec.chapterKeys，不多、不少、不重复。",
        "每章都需要明确触发角色、具体 actionObject、不可逆后果与 progressMode。",
        "每章至少通过 stateWrites、决策、证据开关、resourceDeltas 或失败分支改变局面。",
        "stateReads 只能读取此前已写入或拥有 initialValue 的状态。",
        "有 stateReads 时必须声明 all/any，并提供不同的 onReadPass 与 onReadFail variant。",
        "onReadFail 不能让固定章节消失，必须有 fallbackAction 和资源、状态或证据代价。",
        "除最终章外，本章写入必须被后续章节或结局读取；否则判定“假后果”。",
        "证据开关必须引用真实证据，不能同章同时锁定与解锁同一 key。",
        "推理、政治、生存题材每章需要实质决策；情感、综艺和混合题材至少 60% 章节有决策。",
        "只有 genreMechanic.chapterKeys 中声明的章节才强制实际使用机制，避免所有题材机械化。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.6 累计结局、可达性与冲突裁决", 2)
    for value in [
        "状态变量声明 valueType、initialValue、allowedValues、首次写入章节与意义；资源单独在 resources 登记。",
        "routes.requirements 使用 targetType/targetKey，可读取 state、resource 或 evidence。",
        "五章主要路线至少读取两个不同章节产生的条件：一个来自前半段，一个来自后半段且最终章之前。",
        "至少一个资源、关系或权限状态在不同章节被多次更新，而不是只赋值一次。",
        "至少 60% 的决策状态必须被后续章节或结局读取。",
        "每条条件值必须能由初值、decision、stateWrites、resourceDeltas 或证据开关实际写出。",
        "requirementMode 固定为 all；priority 唯一；冲突由 highest-priority 裁决。",
        "必须恰好有一条无 requirements 的默认路线，defaultRouteKey 必须指向它。",
        "条件完全相同的路线视为重复路线，禁止用最后临时投票替代累计结果。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "5.7 逻辑冲突与模板痕迹", 2)
    add_body(doc, "以下类型属于明确拒绝项：")
    for value in [
        "“真凶或幕后黑手”“凶手或幕后黑手”等未确定责任人。",
        "“实为 A 但又实为 B”式同句互斥。",
        "“待定、尚未确定、任选其一、可能是 A 也可能是 B”。",
        "playerAction、chapterActions.action、activePlan、irreversibleConsequence 或 nextState 大量使用“调查、质问、交换、影响后续”等泛化句。",
        "同篇大量行动只替换角色名或动作对象，字符语义相似度过高。",
        "chapterCausalPattern 仍是“发现—质疑—承认—锁定—投票”的通用五段式。",
        "batchFingerprint 使用“调查旧案、缺失记录、机构阴谋、公开/隐瞒/销毁真相、真相与利益、人性或道德抉择”等泛化短语。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_code_block(
        doc,
        (
            "PLACEHOLDER_NAME = /^(?:角色|玩家|队员|嘉宾|嫌疑人|成员|人物|role|player)...$/iu\n"
            "UNRESOLVED_LOGIC = /真凶...或|凶手或幕后黑手|实为...但又实为|待定|尚未确定|任选其一.../iu\n"
            "GENERIC_FINGERPRINT = /调查旧案|缺失记录|机构阴谋|公开真相|真相与利益|人性抉择|道德抉择.../iu"
        ),
        "关键拒绝模式（便于审阅，完整正则以源码为准）",
    )

    add_heading(doc, "6. 自动修复机制", 1)
    add_body(
        doc,
        "V2.2 不再把所有失败都当成“补字段”。校验器返回 repairMode；局部结构问题走 patch，底层创意或因果失败走 rebuild。"
    )
    repair_decimal_num = make_numbering(doc, "decimal", "%1.", FONT_LATIN)
    for value in [
        "第一次生成：temperature=0.45，允许题材和角色设计具有创造性。",
        "严格校验：收集所有 issues，而不是遇到第一项就停止。",
        "patch：漏字段、key 引用、实体/资源/状态登记或数量错误；携带上一版 JSON，只修改局部结构。",
        "rebuild：高概念降级、贡献结构失败、题材机制不可操作、伪双源、聚光拥挤、泛化行动、不可达结局或批次同构。",
        "rebuild 只携带失败结构的负面摘要，不附完整旧 JSON；从 brief 重做真相、贡献锚点、章节因果和结局。",
        "patch temperature=0.25；rebuild temperature=0.55，避免低温复刻坏地基。",
        "最多尝试三次；第三次仍失败则向上返回错误，不输出半成品。",
    ]:
        add_numbered(doc, value, repair_decimal_num)
    repair_pseudocode = (
        "maxAttempts = clamp(DEEPSEEK_OUTLINE_MAX_ATTEMPTS, 1, 3, default=3)\n"
        "for attempt in 1..maxAttempts:\n"
        "    temperature = 0.45 if initial else (0.25 if patch else 0.55)\n"
        "    result = DeepSeek(messages, maxTokens=12000, temperature=temperature)\n"
        "    record usage, finishReason, completionTokens and nearCompletionLimit\n"
        "    try:\n"
        "        outline = validateStoryOutline(result, spec, strict=true, brief=brief)\n"
        "        return { generationAttempts: attempt, outline }\n"
        "    catch DEEPSEEK_OUTPUT_INVALID as error:\n"
        "        if error.details.repairMode == 'rebuild':\n"
        "            messages = basePrompt + rebuildFromBrief(error.details.issues)\n"
        "        else:\n"
        "            messages = basePrompt + previousJson + targetedPatch(error.details.issues)\n"
        "throw lastError"
    )
    add_code_block(doc, repair_pseudocode)

    add_heading(doc, "6.1 兼容策略", 2)
    add_table(
        doc,
        ["版本", "读取", "严格校验", "允许扩写", "处理方式"],
        [
            ("V1", "允许", "不补猜", "否", "保留旧字段，标记 legacy-outline-v1"),
            ("V2.0", "允许", "不冒充 V2.2", "否", "重新生成或人工迁移"),
            ("V2.1", "允许", "不冒充 V2.2", "否", "补齐实体资源、失败分支与因果路径"),
            ("V2.2", "允许", "全部作用性门禁", "仅通过后允许", "标记 player-driven-outline-v2.2"),
        ],
        [950, 950, 2200, 1800, 3460],
    )
    add_callout(
        doc,
        "兼容边界",
        "旧大纲不会被静默升级成 V2.2，也不会因为仍可显示就被误认为合格。进入扩写前必须补齐题材贡献、实体与资源注册、条件失败 variant、聚光分布和结局因果路径。",
        GOLD,
    )

    add_heading(doc, "7. 四十篇批次生成框架", 1)
    add_body(
        doc,
        "批次脚本先执行每篇 V2.2 严格校验，再比较当前批次与历史库。旧 items 会因 outlineRevision≠2.2 或 readyForExpansion=false 自动失效并重新生成。"
    )
    add_heading(doc, "7.1 单篇批次质量项", 2)
    batch_decimal_num = make_numbering(doc, "decimal", "%1.", FONT_LATIN)
    for value in [
        "题材锚点至少命中一项。",
        "outlineVersion=2、outlineRevision=2.2 且 readyForExpansion=true。",
        "玩家人数与 spec 一致，且贡献锚点、聚光章和结局因果路径完整。",
        "hookPromises 已兑现并引用至少两个题材适配 supportKeys。",
        "实体与资源注册完整；mystery 的 evidenceGraph 才强制核心结论和独立来源。",
        "每章具备稳定 actionTargetKey、失败 variant、不可逆后果和结构化影响。",
        "结局包含可达条件、唯一优先级、冲突规则和默认路线。",
        "十一类 batchFingerprint 全部存在。",
        "禁止回退到预设的通用校园、同学会或埋尸旧案模板。",
    ]:
        add_bullet(doc, value, bullet_num)

    add_heading(doc, "7.2 跨大纲差异度门禁", 2)
    diversity_rows = [
        ("人物姓名", "同名最多出现 2 次", "防止“陈默、林博士、苏晴”等批量复用"),
        ("storyEngine", "同值最多 2 次", "剧情发动机不能只换名词"),
        ("antagonistType", "同值最多 4 次", "机构、管理者、科学家等阻力类型去重"),
        ("finalChoiceType", "同值最多 3 次", "减少“是否公开/销毁/牺牲”的重复"),
        ("themeExpression", "同值最多 2 次", "主题表达不能只是“真相与人性”"),
        ("新增七维", "异常对象、揭示方式、关系拓扑、因果模式、证据组合、权力结构、结局机制", "识别换词后的结构同构"),
        ("单维近似", "字符 bigram Jaccard ≥ 0.78 警告", "这是文本近似，不冒充语义判断"),
        ("组合相似", "十一维等权均值 ≥ 0.72 警告", "默认进入人工复核，不直接误杀"),
        ("校准数据", "目标 100–200 对人工标签", "比较 0.65/0.70/0.72/0.75/0.80 的误杀和漏检"),
        ("历史指纹库", "最多保留 400 条已验收记录", "不只比较当前 40 篇"),
    ]
    add_table(doc, ["维度", "当前门槛", "目的"], diversity_rows, [1850, 2800, 4710])

    add_heading(doc, "7.3 批次失败处理", 2)
    for value in [
        "每篇最多进行脚本级重试；单篇仍失败则记录失败项目。",
        "只要有单篇失败，整批不生成成功汇总。",
        "全部单篇成功后写出 batch-diversity-report.json。",
        "精确重复与频次超限仍可让整批失败；近似阈值校准前只写 warnings。",
        "人工标注完成后可将 OUTLINE_SIMILARITY_ENFORCEMENT 从 review 改为 reject。",
        "记录平均 completion tokens、P95、near-limit 数、JSON 截断数和前后字段长度比。",
        "若 P95≥10,800、近上限比例过高或后半字段显著压缩，报告 recommendTwoStageGeneration=true。",
        "全部通过后才更新历史库，并生成 JSON 与 Markdown 汇总。",
    ]:
        add_numbered(doc, value, batch_decimal_num)

    add_heading(doc, "8. 文件与职责地图", 1)
    file_rows = [
        ("backend/src/prompts/outline.js", "V2.2 系统提示词、完整 schema 与动态章节覆盖公式"),
        ("backend/src/outline-quality-validator.js", "题材贡献、实体资源、失败分支、结局路径与批次相似度"),
        ("backend/src/deepseek-validators.js", "旧版读取与 V2.2 严格准入"),
        ("backend/src/deepseek-client.js", "输出 token、finishReason 与截断识别"),
        ("backend/src/deepseek.js", "生成指标、patch/rebuild 修复分流和下游传递"),
        ("backend/src/prompts/shared.js", "compactOutline 向后续层传递 V2.2 信息"),
        ("examples/pending-review/generate-forty-outlines.mjs", "40 题材、十一维/历史去重、token 压力报告与汇总"),
        ("backend/test/outline-quality-validator.test.js", "伪双源、循环派生、泛化行动、不可达路线、题材门禁与历史同构测试"),
        ("backend/scripts/calibrate-outline-similarity.mjs", "用 100–200 对人工标签评估多档阈值的误杀率与召回"),
        ("backend/scripts/verify-outline-v2-live.mjs", "使用真实 DeepSeek 的 V2.2 UTF-8 冒烟测试"),
    ]
    add_table(doc, ["文件", "职责"], file_rows, [3750, 5610])

    add_heading(doc, "9. 运行与验收说明", 1)
    add_heading(doc, "9.1 环境变量", 2)
    add_code_block(
        doc,
        (
            "DEEPSEEK_API_KEY=<仅保存在服务端环境>\n"
            "DEEPSEEK_BASE_URL=https://api.deepseek.com\n"
            "DEEPSEEK_MODEL=deepseek-v4-flash\n"
            "DEEPSEEK_TIMEOUT_MS=180000\n"
            "DEEPSEEK_OUTLINE_MAX_ATTEMPTS=3\n"
            "OUTLINE_SIMILARITY_ENFORCEMENT=review\n"
            "OUTLINE_FIELD_SIMILARITY_THRESHOLD=0.78\n"
            "OUTLINE_COMPOSITE_SIMILARITY_THRESHOLD=0.72"
        ),
    )
    add_callout(
        doc,
        "安全要求",
        "文档和前端均不包含 API Key。密钥只应存在于 Railway/服务端环境变量或用户 BYOK 加密存储中。",
        RED,
    )

    add_heading(doc, "9.2 验证命令", 2)
    add_code_block(
        doc,
        (
            "# 后端单元测试\n"
            "node --test --test-concurrency=1 test/outline-quality-validator.test.js test/deepseek-pipeline.test.js\n\n"
            "# 模块与依赖图检查\n"
            "npm run check\n\n"
            "# 使用真实 DeepSeek 的 V2.2 冒烟测试（需要服务端密钥）\n"
            "npm run test:outline-v2:live\n\n"
            "# 使用人工标注对校准十一维相似度阈值\n"
            "node scripts/calibrate-outline-similarity.mjs ../examples/pending-review/similarity-calibration-labeled.json\n\n"
            "# 从仓库根目录重新跑 40 篇（需要服务端密钥）\n"
            "node --env-file=.env.railway examples/pending-review/generate-forty-outlines.mjs"
        ),
    )

    add_heading(doc, "9.3 编辑验收清单", 2)
    checklist = [
        "梗概中的每个异常点是否在 hookPromises 中逐项兑现？",
        "玩家贡献锚点是否随题材变化，而不是所有人都被迫提供核心证据？",
        "是否有人被强行每章填入无实质作用的调查动作？",
        "chapterActions 是否满足 ceil(chapterCount×0.6) 的动态覆盖数？",
        "核心责任人是否属于玩家，而非仅由 NPC 承担？",
        "NPC、机构、系统、设备和物证是否都在 entities 中使用稳定 key 登记？",
        "资源初值、上下限、所有者和数字 amount 是否在 resources 中一致？",
        "每个核心结论是否来自不同 provenanceGroup 或不同派生根？",
        "口供、日记、截图和操作记录是否其实来自同一原始主体或系统？",
        "misdirections 是否按嫌疑、记忆、阵营、公众叙事或风险适配题材？",
        "stateReads 不满足时是否进入明确 fallback variant，并支付结构化代价？",
        "每章 stateWrites、证据开关或资源变化是否真的被下游读取？",
        "题材是否使用适配的 progressMode，而不是统一塞物证和二选一？",
        "结局路线是否可达、条件为 AND、优先级唯一且有默认路线？",
        "至少 60% 的决策状态是否被后续章节或结局读取？",
        "最终章是否仍依赖突然出现的完整答案文件？",
        "聚光章是否分布合理，并与 turnChapterKeys 和 triggerRoleKeys 对齐？",
        "每名玩家的状态、资源或证据影响是否存在通往结局条件的路径？",
        "十一维字符相似度警告是否经过人工复核，而非被误称为语义定论？",
    ]
    for value in checklist:
        paragraph = add_bullet(doc, f"□ {value}", bullet_num)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, size=9.5, color=BLACK)

    add_heading(doc, "10. 当前边界与下一步", 1)
    add_body(
        doc,
        "V2.2 已从“字段是否存在”推进到“字段是否真正生效”，但仍不能代替资深编辑判断惊奇强度、审美、动机、史实、专业流程和对白。机器通过只代表具备可验证的扩写地基，不代表达到发行品质。"
    )
    add_callout(
        doc,
        "建议下一轮",
        "先完成 100–200 对相似度人工标注并校准阈值，再用 V2.2 重跑四十篇；同时检查作用性失败、token P95、首尾字段长度和 8–12 篇人工深审结果。",
        GREEN,
    )
    for value in [
        "把人工审稿结论继续转化为可机器验证的字段或状态引用，而不是只增加提示词形容词。",
        "分别统计 patch 与 rebuild 的通过率、平均修复次数和常见失败项，判断哪些问题来自漏字段，哪些来自坏地基。",
        "进入完整角色本前增加实体—资源—状态—结局因果图审查和小规模 AI/真人跑团。",
        "历史指纹库保留最近 400 条；字符 bigram 校准后仍可接入词向量或第二模型做成对判断。",
    ]:
        paragraph = add_bullet(doc, value, bullet_num)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, size=8.5, color=BLACK)

    # Final preset audit notes are encoded in styles and geometry; save.
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_document()
